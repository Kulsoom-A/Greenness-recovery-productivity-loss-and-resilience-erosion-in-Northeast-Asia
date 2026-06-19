from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


MAIN = Path(r"F:\HarbinPaper-02-\Manuscript\Paper2_Main_Manuscript_Final.docx")

REFERENCES = [
    "Abatzoglou, J. T., Dobrowski, S. Z., Parks, S. A., & Hegewisch, K. C. (2018). TerraClimate, a high-resolution global dataset of monthly climate and climatic water balance from 1958-2015. Scientific Data, 5, 170191.",
    "Allen, C. D., Macalady, A. K., Chenchouni, H., et al. (2010). A global overview of drought and heat-induced tree mortality reveals emerging climate change risks for forests. Forest Ecology and Management, 259(4), 660-684.",
    "Anderegg, W. R. L., Schwalm, C., Biondi, F., et al. (2015). Pervasive drought legacies in forest ecosystems and their implications for carbon cycle models. Science, 349(6247), 528-532.",
    "Anderegg, W. R. L., Trugman, A. T., Badgley, G., et al. (2020). Climate-driven risks to the climate mitigation potential of forests. Science, 368(6497), eaaz7005.",
    "Badgley, G., Field, C. B., & Berry, J. A. (2017). Canopy near-infrared reflectance and terrestrial photosynthesis. Science Advances, 3(3), e1602244.",
    "Camps-Valls, G., Campos-Taberner, M., Moreno-Martínez, Á., et al. (2021). A unified vegetation index for quantifying the terrestrial biosphere. Science Advances, 7(9), eabc7447.",
    "Chen, C., Park, T., Wang, X., et al. (2019). China and India lead in greening of the world through land-use management. Nature Sustainability, 2(2), 122-129.",
    "Cheng, R. (2024). Solar-induced chlorophyll fluorescence (SIF): Towards a better understanding of vegetation dynamics and carbon uptake in Arctic-boreal ecosystems. Current Climate Change Reports, 10, 13-32.",
    "DeSoto, L., Cailleret, M., Sterck, F., et al. (2020). Low growth resilience to drought is related to future mortality risk in trees. Nature Communications, 11, 545.",
    "Forzieri, G., Dakos, V., McDowell, N. G., Ramdane, A., & Cescatti, A. (2022). Emerging signals of declining forest resilience under climate change. Nature, 608(7923), 534-539.",
    "Frank, D., Reichstein, M., Bahn, M., et al. (2015). Effects of climate extremes on the terrestrial carbon cycle: concepts, processes and potential future impacts. Global Change Biology, 21(8), 2861-2880.",
    "Giglio, L., Boschetti, L., Roy, D. P., Humber, M. L., & Justice, C. O. (2018). The Collection 6 MODIS burned area mapping algorithm and product. Remote Sensing of Environment, 217, 72-85.",
    "Hammond, W. M., Williams, A. P., Abatzoglou, J. T., et al. (2022). Global field observations of tree die-off reveal hotter-drought fingerprint for Earth's forests. Nature Communications, 13, 1761.",
    "Hansen, M. C., Potapov, P. V., Moore, R., et al. (2013). High-resolution global maps of 21st-century forest cover change. Science, 342(6160), 850-853.",
    "Huete, A., Didan, K., Miura, T., Rodriguez, E. P., Gao, X., & Ferreira, L. G. (2002). Overview of the radiometric and biophysical performance of the MODIS vegetation indices. Remote Sensing of Environment, 83(1-2), 195-213.",
    "Li, X., & Xiao, J. (2019). A global, 0.05-degree product of solar-induced chlorophyll fluorescence derived from OCO-2, MODIS, and reanalysis data. Remote Sensing, 11(5), 517.",
    "McDowell, N., Pockman, W. T., Allen, C. D., et al. (2008). Mechanisms of plant survival and mortality during drought: why do some plants survive while others succumb to drought? New Phytologist, 178(4), 719-739.",
    "McDowell, N. G., Allen, C. D., Anderson-Teixeira, K., et al. (2020). Pervasive shifts in forest dynamics in a changing world. Science, 368(6494), eaaz9463.",
    "Pan, Y., Birdsey, R. A., Fang, J., et al. (2011). A large and persistent carbon sink in the world's forests. Science, 333(6045), 988-993.",
    "Reichstein, M., Bahn, M., Ciais, P., et al. (2013). Climate extremes and the carbon cycle. Nature, 500(7462), 287-295.",
    "Ruehr, N. K., Grote, R., Mayr, S., & Arneth, A. (2019). Beyond the extreme: recovery of carbon and water relations in woody plants following heat and drought stress. Tree Physiology, 39(8), 1285-1299.",
    "Running, S. W., Nemani, R. R., Heinsch, F. A., et al. (2004). A continuous satellite-derived measure of global terrestrial primary production. BioScience, 54(6), 547-560.",
    "Schwalm, C. R., Anderegg, W. R. L., Michalak, A. M., et al. (2017). Global patterns of drought recovery. Nature, 548(7666), 202-205.",
    "Stocker, B. D., Zscheischler, J., Keenan, T. F., et al. (2019). Drought impacts on terrestrial primary production underestimated by satellite monitoring. Nature Geoscience, 12(4), 264-270.",
    "Sun, Y., Frankenberg, C., Wood, J. D., et al. (2017). OCO-2 advances photosynthesis observation from space via solar-induced chlorophyll fluorescence. Science, 358(6360), eaam5747.",
    "Tang, Z., Miralles, D. G., Guo, Z., & Maes, W. H. (2026). Fast response of satellite fluorescence-derived plant physiology to drought stress. Nature Communications, 17, 2886.",
    "Vicente-Serrano, S. M., Beguería, S., & López-Moreno, J. I. (2010). A multiscalar drought index sensitive to global warming: the standardized precipitation evapotranspiration index. Journal of Climate, 23(7), 1696-1718.",
    "Zhao, M., & Running, S. W. (2010). Drought-induced reduction in global terrestrial net primary production from 2000 through 2009. Science, 329(5994), 940-943.",
    "Zhu, Z., Piao, S., Myneni, R. B., et al. (2016). Greening of the Earth and its drivers. Nature Climate Change, 6(8), 791-795.",
    "Zscheischler, J., Westra, S., van den Hurk, B. J. J. M., et al. (2018). Future climate risk from compound events. Nature Climate Change, 8(6), 469-477.",
    "Zscheischler, J., Martius, O., Westra, S., et al. (2020). A typology of compound weather and climate events. Nature Reviews Earth & Environment, 1(7), 333-347.",
]


CONCLUSION = (
    "This study shows that apparent vegetation recovery after compound hot-dry events can conceal a longer and more consequential functional legacy. "
    "Across Northeast Asia, greenness metrics often returned toward pre-event conditions while MOD17 GPP and an independent fluorescence proxy continued to indicate suppressed productivity. "
    "The persistence of this mismatch across kNDVI, EVI, NIRv, MOD17 GPP, and GOSIF demonstrates that the signal is not reducible to a single spectral index or productivity algorithm, although the weaker SIF response appropriately tempers the magnitude of MOD17-based carbon-loss estimates. "
    "Repeated exposure further intensified productivity loss, indicating resilience erosion that was not explained by baseline contamination or by pixel-level uncertainty. "
    "The strongest vulnerability was concentrated in deciduous broadleaf systems and in areas where exposure, productivity loss, recurrence sensitivity, and hidden mismatch coincided. "
    "Together, these findings argue that post-extreme recovery assessments should not rely on greenness alone. "
    "A robust monitoring framework for climate-stressed ecosystems should combine structural greenness, productivity or fluorescence, event-history metrics, and spatially explicit uncertainty tests so that apparent canopy recovery is not mistaken for restored carbon uptake."
)


def set_run(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic


def format_p(p, heading=False, refs=False):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run(run, size=12, bold=heading)


def replace_text(p, old, new):
    if old not in p.text:
        return False
    text = p.text.replace(old, new)
    p.clear()
    p.add_run(text)
    return True


def remove_from_first_references(doc: Document) -> None:
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().lower() == "references":
            start = i
            break
    if start is None:
        return
    for p in list(doc.paragraphs[start:]):
        p._element.getparent().remove(p._element)


def append_section(doc: Document, title: str, body: str | None = None):
    h = doc.add_paragraph()
    h.add_run(title)
    format_p(h, heading=True)
    if body:
        p = doc.add_paragraph()
        p.add_run(body)
        format_p(p)


def main() -> None:
    doc = Document(MAIN)

    for p in doc.paragraphs:
        replace_text(
            p,
            "Annual kNDVI, EVI, NIRv, monthly SPEI-12, TerraClimate maximum temperature, precipitation, potential evapotranspiration, climatic water deficit, PDSI, annual MOD17 GPP, annual GOSIF solar-induced fluorescence, and the six-class vegetation mask were harmonized to a common 0.1 degree EPSG:4326 grid.",
            "Annual kNDVI, EVI, NIRv, monthly SPEI-12, TerraClimate maximum temperature, precipitation, potential evapotranspiration, climatic water deficit, PDSI, annual MOD17 GPP, annual GOSIF solar-induced fluorescence, and the six-class vegetation mask were harmonized to a common 0.1 degree EPSG:4326 grid (Camps-Valls et al., 2021; Huete et al., 2002; Vicente-Serrano et al., 2010; Abatzoglou et al., 2018; Running et al., 2004; Li & Xiao, 2019).",
        )
        replace_text(
            p,
            "MODIS burned area and Hansen forest-loss products were used as external disturbance screens.",
            "MODIS burned area and Hansen forest-loss products were used as external disturbance screens (Giglio et al., 2018; Hansen et al., 2013).",
        )
        replace_text(
            p,
            "Signals that track photosynthesis more directly gross primary productivity and solar-induced chlorophyll fluorescence, the latter emitted during photochemistry itself can diverge from greenness when function is impaired while canopy structure remains intact (Sun et al., 2017; Stocker et al., 2019).",
            "Signals that track photosynthesis more directly, including gross primary productivity and solar-induced chlorophyll fluorescence, can diverge from greenness when function is impaired while canopy structure remains intact (Sun et al., 2017; Stocker et al., 2019; Tang et al., 2026).",
        )
        replace_text(
            p,
            "Because fluorescence is emitted during photochemistry and tracks photosynthetic activity more directly than reflectance, its agreement with the GPP-based signal points to a functional decoupling rather than to the behaviour of a single algorithm.",
            "Because fluorescence is emitted during photochemistry and tracks photosynthetic activity more directly than reflectance, its agreement with the GPP-based signal points to a functional decoupling rather than to the behaviour of a single algorithm (Cheng, 2024; Tang et al., 2026).",
        )
        if p.text.strip().startswith("The third is for risk and management.") and "Hammond et al., 2022" not in p.text:
            replace_text(
                p,
                "(Anderegg et al., 2020; Hammond et al., 2022)",
                "(Anderegg et al., 2020; Hammond et al., 2022)",
            )

    remove_from_first_references(doc)
    append_section(doc, "Conclusion", CONCLUSION)
    append_section(doc, "References")
    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.add_run(ref)
        format_p(p)

    for p in doc.paragraphs:
        heading = p.text.strip() in {
            "Abstract", "Introduction", "Materials and methods", "Study area", "Datasets and preprocessing",
            "Compound hot-dry events and response metrics", "Recurrent exposure, mismatch, and hotspot metrics",
            "Driver modelling and robustness checks", "External disturbance control", "Results",
            "Event legacies and recurrent exposure", "Greenness-productivity mismatch and integrated hotspots",
            "Spatial controls, model validation, and robustness", "Discussion", "Principal findings",
            "Greenness recovery can mask incomplete functional recovery", "A physiological basis for the lag",
            "A multi-proxy framework for separating apparent from functional recovery",
            "MOD17 amplifies, but does not manufacture, the signal", "Recurrent exposure and the erosion of resilience",
            "The compound framing, read honestly", "Vegetation-type and regional structure",
            "Implications for monitoring, carbon assessment, and risk", "Limitations", "Conclusion", "References",
        }
        format_p(p, heading=heading)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    format_p(p)

    doc.save(MAIN)


if __name__ == "__main__":
    main()
