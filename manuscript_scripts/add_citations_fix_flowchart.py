from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


MAIN = Path(r"F:\HarbinPaper-02-\Manuscript\Paper2_Main_Manuscript_Final.docx")


def set_run(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic


def format_p(p, heading=False, italic=False):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run(run, size=12, bold=heading, italic=italic)


def replace_text(p, old, new):
    if old not in p.text:
        return False
    text = p.text.replace(old, new)
    p.clear()
    p.add_run(text)
    return True


def insert_before(paragraph, text, italic=False):
    new_p = paragraph.insert_paragraph_before(text)
    format_p(new_p, italic=italic)
    return new_p


def remove_duplicate_flowchart_caption(doc: Document):
    caption = "Figure 2. Methodological workflow for detecting compound hot-dry events"
    seen = False
    for idx, p in enumerate(list(doc.paragraphs)):
        if p.text.startswith(caption):
            if not seen:
                seen = True
                prior = doc.paragraphs[max(0, idx - 3):idx]
                if not any("Insert final Figure 2" in q.text for q in prior):
                    insert_before(p, "[Insert final Figure 2 methodological workflow / flowchart here]", italic=True)
            else:
                p._element.getparent().remove(p._element)


def add_contextual_citations(doc: Document):
    replacements = {
        "Northeast Asia was selected as the study domain because it contains one of the largest climatically transitional vegetation regions of the Northern Hemisphere, linking boreal and cold-temperate forests, montane forest belts, shrublands, and dry grassland margins within a single heat- and water-stress gradient.": 
        "Northeast Asia was selected as the study domain because it contains one of the largest climatically transitional vegetation regions of the Northern Hemisphere, linking boreal and cold-temperate forests, montane forest belts, shrublands, and dry grassland margins within a single heat- and water-stress gradient (McDowell et al., 2020; Forzieri et al., 2022).",
        "Compound hot–dry events were identified independently for each pixel and year using December SPEI-12 and annual maximum temperature.": 
        "Compound hot–dry events were identified independently for each pixel and year using December SPEI-12 and annual maximum temperature, following the compound-event logic used for co-occurring drought and heat extremes (Vicente-Serrano et al., 2010; Zscheischler et al., 2018, 2020).",
        "To isolate the compound-event interpretation, two mutually exclusive univariate event classes were also evaluated using identical response windows.": 
        "To isolate the compound-event interpretation, two mutually exclusive univariate event classes were also evaluated using identical response windows, consistent with recommendations that compound hazards be compared with their marginal drivers where possible (Zscheischler et al., 2018, 2020).",
        "Because cumulative loss and legacy duration are one-sided metrics, they may produce positive values even under baseline-centered interannual variability.": 
        "Because cumulative loss and legacy duration are one-sided metrics, they may produce positive values even under baseline-centered interannual variability; null-event timing tests were therefore used to distinguish event-aligned legacies from background variability (Anderegg et al., 2015; Schwalm et al., 2017).",
        "To examine whether repeated climate-stress exposure was associated with resilience erosion, valid events were ranked chronologically within each pixel.": 
        "To examine whether repeated climate-stress exposure was associated with resilience erosion, valid events were ranked chronologically within each pixel, motivated by evidence that repeated drought exposure can reduce growth resilience and increase later vulnerability (DeSoto et al., 2020; Forzieri et al., 2022).",
        "Hidden productivity-loss mismatch was designed to identify cases where greenness recovery suggested apparent ecosystem recovery, but productivity remained suppressed.": 
        "Hidden productivity-loss mismatch was designed to identify cases where greenness recovery suggested apparent ecosystem recovery, but productivity remained suppressed, a risk highlighted by evidence that greenness can underestimate drought impacts on productivity (Stocker et al., 2019; Tang et al., 2026).",
        "A standardized mismatch-intensity metric was also calculated to quantify the divergence between greenness and productivity responses:": 
        "A standardized mismatch-intensity metric was also calculated to quantify the divergence between greenness and productivity responses, following the broader need to separate structural vegetation indices from functional productivity proxies (Badgley et al., 2017; Camps-Valls et al., 2021; Cheng, 2024):",
        "Potential climatic and ecological controls were evaluated using event-window driver layers.": 
        "Potential climatic and ecological controls were evaluated using event-window driver layers derived from SPEI-12, TerraClimate climate and water-balance variables, pre-event productivity, event frequency, and vegetation class (Vicente-Serrano et al., 2010; Abatzoglou et al., 2018).",
        "External disturbance was screened using MODIS MCD64A1 burned area and Hansen Global Forest Change tree-cover loss products downloaded from Google Earth Engine and aligned to the 0.1 degree analysis grid.": 
        "External disturbance was screened using MODIS MCD64A1 burned area and Hansen Global Forest Change tree-cover loss products downloaded from Google Earth Engine and aligned to the 0.1 degree analysis grid (Giglio et al., 2018; Hansen et al., 2013).",
        "The independent GOSIF validation directly addressed whether the hidden productivity-loss result depended on the MOD17 GPP light-use-efficiency algorithm.": 
        "The independent GOSIF validation directly addressed whether the hidden productivity-loss result depended on the MOD17 GPP light-use-efficiency algorithm, because SIF provides an independent photosynthetic constraint on productivity dynamics (Running et al., 2004; Li & Xiao, 2019; Cheng, 2024).",
        "The greenness-side robustness test further showed that the SIF-supported mismatch was not unique to the kNDVI transformation.": 
        "The greenness-side robustness test further showed that the SIF-supported mismatch was not unique to the kNDVI transformation, because EVI and NIRv represent structurally different greenness/canopy signals from kNDVI (Huete et al., 2002; Badgley et al., 2017; Camps-Valls et al., 2021).",
        "This study set out to test whether the recovery of canopy greenness after compound hot–dry events correspond to recovered ecosystem function across the Northeast Asian vegetation domain.": 
        "This study set out to test whether the recovery of canopy greenness after compound hot–dry events correspond to recovered ecosystem function across the Northeast Asian vegetation domain, a question motivated by known drought legacies and by the documented divergence between structural and functional remote-sensing signals (Anderegg et al., 2015; Stocker et al., 2019; Tang et al., 2026).",
        "A methodological contribution of this work is the deliberate separation of the greenness and productivity sides of the mismatch onto independent measurement principles, paired with explicit tests against chance and against spatial structure.": 
        "A methodological contribution of this work is the deliberate separation of the greenness and productivity sides of the mismatch onto independent measurement principles, paired with explicit tests against chance and against spatial structure (Badgley et al., 2017; Li & Xiao, 2019; Camps-Valls et al., 2021; Cheng, 2024).",
        "The first implication is for monitoring. Greenness-based recovery metrics overstate ecosystem recovery after hot–dry extremes and should be paired with fluorescence- or flux-based measures wherever carbon outcomes are the question; relying on greenness alone will systematically misjudge how long ecosystems remain functionally impaired.": 
        "The first implication is for monitoring. Greenness-based recovery metrics can overstate ecosystem recovery after hot–dry extremes and should be paired with fluorescence- or flux-based measures wherever carbon outcomes are the question; relying on greenness alone risks misjudging how long ecosystems remain functionally impaired (Stocker et al., 2019; Cheng, 2024; Tang et al., 2026).",
        "This study shows that apparent vegetation recovery after compound hot-dry events can conceal a longer and more consequential functional legacy.": 
        "This study shows that apparent vegetation recovery after compound hot-dry events can conceal a longer and more consequential functional legacy, consistent with earlier evidence for drought legacies in ecosystem carbon uptake and growth (Anderegg et al., 2015; Schwalm et al., 2017).",
        "where  represents the relative year from to , and denotes the event year.": 
        "where  represents the relative year from to , and denotes the event year. This anomaly-window design follows the broader drought-legacy framing used to quantify delayed vegetation recovery after climate extremes (Anderegg et al., 2015; Schwalm et al., 2017).",
        "Cumulative loss was defined as the sum of negative anomalies from the event year to four years after the event:": 
        "Cumulative loss was defined as the sum of negative anomalies from the event year to four years after the event, consistent with the interpretation of post-event carbon suppression as an accumulated legacy effect (Reichstein et al., 2013; Frank et al., 2015):",
        "This metric represents the cumulative suppression of vegetation activity relative to the pre-event baseline. Larger positive values indicate greater post-event vegetation loss. Legacy duration was defined as the number of post-event years during which the response variable remained below the pre-event baseline:": 
        "This metric represents the cumulative suppression of vegetation activity relative to the pre-event baseline. Larger positive values indicate greater post-event vegetation loss. Legacy duration was defined as the number of post-event years during which the response variable remained below the pre-event baseline, following drought-recovery studies that treat delayed return to baseline as a legacy signal (Anderegg et al., 2015; Schwalm et al., 2017):",
        "where is an indicator function equal to 1 when the condition is true and 0 otherwise. Therefore, ranges from 0 to 4, with higher values indicating longer post-event legacy effects.": 
        "where is an indicator function equal to 1 when the condition is true and 0 otherwise. Therefore, ranges from 0 to 4, with higher values indicating longer post-event legacy effects (Anderegg et al., 2015; Schwalm et al., 2017).",
        "where represents resistance at pixel during event year , and is the anomaly in the event year. More negative values indicate weaker resistance and stronger immediate vegetation suppression.": 
        "where represents resistance at pixel during event year , and is the anomaly in the event year. More negative values indicate weaker resistance and stronger immediate vegetation suppression, consistent with resilience concepts based on resistance and recovery after drought (DeSoto et al., 2020; Forzieri et al., 2022).",
        "Recovery time was defined as the first post-event year in which the anomaly returned to or exceeded the pre-event baseline:": 
        "Recovery time was defined as the first post-event year in which the anomaly returned to or exceeded the pre-event baseline, a common operational definition of return-to-baseline recovery after drought disturbance (Schwalm et al., 2017; Ruehr et al., 2019):",
        "Thus, larger values of indicate slower recovery, while denotes an unrecovered event within the observation window.": 
        "Thus, larger values of indicate slower recovery, while denotes an unrecovered event within the observation window (Schwalm et al., 2017; Ruehr et al., 2019).",
        "For each pixel, later-minus-first differences were calculated for resistance, recovery time, and cumulative loss. For a given response metric , the recurrent-event change was calculated as:": 
        "For each pixel, later-minus-first differences were calculated for resistance, recovery time, and cumulative loss. For a given response metric , the recurrent-event change was calculated as a direct test of whether repeated drought/heat exposure erodes recovery capacity (DeSoto et al., 2020; Forzieri et al., 2022):",
        "Mismatch duration was calculated as the total number of post-event years with hidden productivity suppression across all valid events:": 
        "Mismatch duration was calculated as the total number of post-event years with hidden productivity suppression across all valid events, reflecting the need to distinguish structural greenness recovery from functional carbon recovery (Stocker et al., 2019; Tang et al., 2026):",
        "where is an indicator function equal to 1 when the condition is true and 0 otherwise.": 
        "where is an indicator function equal to 1 when the condition is true and 0 otherwise.",
        "Mismatch prevalence was calculated as the fraction of event-affected pixels showing at least one hidden mismatch:": 
        "Mismatch prevalence was calculated as the fraction of event-affected pixels showing at least one hidden mismatch, providing a spatial prevalence estimate of greenness-productivity decoupling (Stocker et al., 2019; Cheng, 2024):",
        "where is the number of event-affected pixels with mismatch duration greater than zero, and is the total number of event-affected pixels.": 
        "where is the number of event-affected pixels with mismatch duration greater than zero, and is the total number of event-affected pixels.",
        "where denotes standardized anomalies across event-affected forest pixels. Positive values of indicate stronger apparent recovery in kNDVI than in GPP, suggesting that greenness recovery may mask continued productivity suppression. This metric was averaged across post-event years to represent overall post-event kNDVI–GPP mismatch intensity. A vulnerability-hotspot classification was then derived by integrating exposure, productivity loss, legacy duration, recurrent-event change, and kNDVI–GPP mismatch metrics. Each continuous component was scaled from 0 to 1 using the 5th and 95th percentiles across event-affected forest pixels:": 
        "where denotes standardized anomalies across event-affected forest pixels. Positive values of indicate stronger apparent recovery in kNDVI than in GPP, suggesting that greenness recovery may mask continued productivity suppression. This metric was averaged across post-event years to represent overall post-event kNDVI-GPP mismatch intensity. A vulnerability-hotspot classification was then derived by integrating exposure, productivity loss, legacy duration, recurrent-event change, and kNDVI-GPP mismatch metrics, consistent with multi-metric resilience assessments that combine exposure, resistance, and recovery information (DeSoto et al., 2020; Forzieri et al., 2022). Each continuous component was scaled from 0 to 1 using the 5th and 95th percentiles across event-affected forest pixels:",
        "where is the original value of metric at pixel , and and are the 5th and 95th percentile values of that metric. Values below 0 were set to 0, and values above 1 were set to 1. The composite vulnerability index was calculated as the mean of the scaled components:": 
        "where is the original value of metric at pixel , and and are the 5th and 95th percentile values of that metric. Values below 0 were set to 0, and values above 1 were set to 1. The composite vulnerability index was calculated as the mean of the scaled components, and was interpreted as a descriptive climate-response synthesis rather than a mortality predictor (Forzieri et al., 2022; Hansen et al., 2013):",
        "where is the composite vulnerability index at pixel , is the scaled value of component , and is the total number of components. The components included event frequency, GPP cumulative loss, GPP legacy duration, later-minus-first GPP cumulative-loss change, later-minus-first GPP recovery-time change, hidden mismatch duration, and standardized mismatch intensity. Pixels were then assigned to ordered vulnerability classes representing low-impact recovery, delayed productivity recovery, hidden productivity suppression, recurrent resilience erosion, and severe multi-metric hotspots. This classification identified forest areas where repeated exposure, persistent productivity loss, delayed recovery, and kNDVI–GPP mismatch jointly indicated elevated ecological vulnerability.": 
        "where is the composite vulnerability index at pixel , is the scaled value of component , and is the total number of components. The components included event frequency, GPP cumulative loss, GPP legacy duration, later-minus-first GPP cumulative-loss change, later-minus-first GPP recovery-time change, hidden mismatch duration, and standardized mismatch intensity. Pixels were then assigned to ordered vulnerability classes representing low-impact recovery, delayed productivity recovery, hidden productivity suppression, recurrent resilience erosion, and severe multi-metric hotspots. This classification identified forest areas where repeated exposure, persistent productivity loss, delayed recovery, and kNDVI-GPP mismatch jointly indicated elevated ecological vulnerability (DeSoto et al., 2020; Forzieri et al., 2022).",
    }
    for p in doc.paragraphs:
        for old, new in replacements.items():
            replace_text(p, old, new)


def main():
    doc = Document(MAIN)
    remove_duplicate_flowchart_caption(doc)
    add_contextual_citations(doc)
    headings = {
        "Abstract", "Introduction", "Materials and methods", "Study area", "Datasets and preprocessing",
        "Compound hot-dry events and response metrics", "Compound hot–dry events and response metrics",
        "Recurrent exposure, mismatch, and hotspot metrics", "Driver modelling and robustness checks",
        "External disturbance control", "Results", "Event legacies and recurrent exposure",
        "Greenness-productivity mismatch and integrated hotspots", "Spatial controls, model validation, and robustness",
        "Discussion", "Principal findings", "Greenness recovery can mask incomplete functional recovery",
        "A physiological basis for the lag", "A multi-proxy framework for separating apparent from functional recovery",
        "MOD17 amplifies, but does not manufacture, the signal", "Recurrent exposure and the erosion of resilience",
        "The compound framing, read honestly", "Vegetation-type and regional structure",
        "Implications for monitoring, carbon assessment, and risk", "Limitations", "Conclusion", "References",
    }
    for p in doc.paragraphs:
        format_p(p, heading=p.text.strip() in headings, italic=p.text.strip().startswith("Figure ") or p.text.strip().startswith("Table ") or p.text.strip().startswith("[Insert"))
    doc.save(MAIN)


if __name__ == "__main__":
    main()
