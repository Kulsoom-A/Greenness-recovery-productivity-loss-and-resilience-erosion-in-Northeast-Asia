from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph


ROOT = Path(r"F:\HarbinPaper-02-\Manuscript")
MAIN = ROOT / "Paper2_Main_Manuscript_Final.docx"
SUPP = ROOT / "Paper2_Supplementary_Material_Final.docx"


def set_run_format(run, size=None, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(0, 0, 0)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_format(p, is_heading=False, is_title=False, after_references=False):
    if after_references:
        # Keep reference text content and paragraph layout as the user supplied it.
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(0, 0, 0)
        return
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if is_title:
        for run in p.runs:
            set_run_format(run, size=14, bold=True)
    elif is_heading:
        for run in p.runs:
            set_run_format(run, size=12, bold=True)
    else:
        for run in p.runs:
            set_run_format(run, size=12)


def paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    out = Paragraph(new_p, paragraph._parent)
    if text:
        out.add_run(text)
    return out


def iter_all_paragraphs(doc: Document):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def replace_paragraph_text_preserve_first_style(p, text):
    if p.runs:
        first = p.runs[0]
        first.text = text
        for run in p.runs[1:]:
            run.text = ""
    else:
        p.add_run(text)


def renumber_main_figures(doc: Document) -> None:
    # Shift every existing main Figure 2-11 reference/caption to Figure 3-12.
    pattern = re.compile(r"\bFigure\s+(1[01]|[2-9])\b")
    for p in iter_all_paragraphs(doc):
        text = p.text
        if not text:
            continue
        def repl(match):
            return f"Figure {int(match.group(1)) + 1}"
        new_text = pattern.sub(repl, text)
        if new_text != text:
            replace_paragraph_text_preserve_first_style(p, new_text)


def insert_flowchart_placeholder(doc: Document) -> None:
    if any("Figure 2. Methodological workflow" in p.text for p in doc.paragraphs):
        return
    table_caption = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("Table 1. Datasets and preprocessing"):
            table_caption = p
            break
    if table_caption is None:
        raise RuntimeError("Could not find Table 1 caption for flowchart insertion.")
    placeholder = paragraph_after(table_caption, "[Insert final Figure 2 methodological workflow / flowchart here]")
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    placeholder.paragraph_format.line_spacing = 1.5
    r = placeholder.runs[0]
    set_run_format(r, size=12, italic=True)
    caption = paragraph_after(
        placeholder,
        "Figure 2. Methodological workflow for detecting compound hot-dry events, estimating greenness and productivity legacies, identifying recurrent-event resilience erosion and hidden productivity mismatch, mapping vulnerability hotspots, and applying robustness and validation checks.",
    )
    caption.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    caption.paragraph_format.line_spacing = 1.5
    for run in caption.runs:
        set_run_format(run, size=12, italic=True)


def add_supplementary_inventory_citation(doc: Document) -> None:
    marker = "Supplementary material inventory"
    if any(marker in p.text for p in doc.paragraphs):
        return
    target = None
    for p in doc.paragraphs:
        if "The datasets, periods, analytical roles, and preprocessing steps are summarized in Table 1." in p.text:
            target = p
            break
    if target is None:
        return
    p = paragraph_after(
        target,
        "Supplementary material inventory: all supporting spatial panels and diagnostic plots are provided in Supplementary Figs. S1-S14, and the corresponding raster inventory, class summaries, uncertainty estimates, model diagnostics, robustness checks, validation tests, and moved legacy table are provided in Supplementary Tables S1-S26.",
    )
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        set_run_format(run, size=12)


def apply_document_format(doc: Document) -> None:
    for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style.font.color.rgb = RGBColor(0, 0, 0)
            if style_name.startswith("Heading"):
                style.font.bold = True
                style.font.size = Pt(12)
            elif style_name == "Title":
                style.font.size = Pt(14)
                style.font.bold = True
            else:
                style.font.size = Pt(12)
    after_refs = False
    for i, p in enumerate(doc.paragraphs):
        stripped = p.text.strip().lower()
        is_refs_heading = stripped in {"references", "reference"}
        is_heading = bool(stripped) and (
            p.style.name.startswith("Heading")
            or stripped in {
                "abstract",
                "introduction",
                "materials and methods",
                "study area",
                "datasets and preprocessing",
                "compound hot-dry events and response metrics",
                "recurrent exposure, mismatch, and hotspot metrics",
                "driver modelling and robustness checks",
                "external disturbance control",
                "results",
                "event legacies and recurrent exposure",
                "greenness-productivity mismatch and integrated hotspots",
                "spatial controls, model validation, and robustness",
                "discussion",
                "conclusion",
                "references",
            }
        )
        set_paragraph_format(p, is_heading=is_heading, is_title=(i == 0), after_references=after_refs)
        if is_refs_heading:
            after_refs = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.line_spacing = 1.5
                    for run in p.runs:
                        set_run_format(run, size=10)


def main() -> None:
    main_doc = Document(MAIN)
    main_text = "\n".join(p.text for p in main_doc.paragraphs)
    if "Figure 12." not in main_text:
        renumber_main_figures(main_doc)
    insert_flowchart_placeholder(main_doc)
    add_supplementary_inventory_citation(main_doc)
    apply_document_format(main_doc)
    main_doc.save(MAIN)

    supp_doc = Document(SUPP)
    apply_document_format(supp_doc)
    supp_doc.save(SUPP)


if __name__ == "__main__":
    main()
