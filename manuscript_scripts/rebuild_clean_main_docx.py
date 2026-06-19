from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"F:\HarbinPaper-02-\Manuscript")
SRC = ROOT / "Paper2_Main_Manuscript_Final.docx"
OUT = ROOT / "Paper2_Main_Manuscript_Final_clean.docx"
PUBFIG = ROOT / "figures_publication"

FIGURE_IMAGES = {
    "Figure 3.": PUBFIG / "figure1_event_panel_ggstyle.png",
    "Figure 4.": PUBFIG / "qgis_guide_figure2_recurrent_resilience_erosion.png",
    "Figure 5.": PUBFIG / "qgis_guide_figure3_hidden_productivity_mismatch.png",
    "Figure 6.": PUBFIG / "qgis_guide_figure4_integrated_vulnerability_hotspots.png",
    "Figure 7.": PUBFIG / "qgis_guide_figure5_spatial_controls_predicted_vulnerability.png",
    "Figure 8.": PUBFIG / "figure11_driver_importance_combined.png",
    "Figure 9.": PUBFIG / "figure12_internal_robustness_checks.png",
    "Figure 10.": PUBFIG / "figure14_sif_productivity_validation.png",
    "Figure 11.": PUBFIG / "figure15_structural_greenness_sif_validation.png",
    "Figure 12.": PUBFIG / "figure13_external_disturbance_control.png",
}

HEADING_TEXT = {
    "Abstract",
    "Introduction",
    "Materials and methods",
    "Study area",
    "Datasets and preprocessing",
    "Compound hot-dry events and response metrics",
    "Recurrent exposure, mismatch, and hotspot metrics",
    "Driver modelling and robustness checks",
    "External disturbance control",
    "Results",
    "Event legacies and recurrent exposure",
    "Greenness-productivity mismatch and integrated hotspots",
    "Spatial controls, model validation, and robustness",
    "Discussion",
    "Principal findings",
    "Greenness recovery can mask incomplete functional recovery",
    "A physiological basis for the lag",
    "A multi-proxy framework for separating apparent from functional recovery",
    "MOD17 amplifies, but does not manufacture, the signal",
    "Recurrent exposure and the erosion of resilience",
    "The compound framing, read honestly",
    "Vegetation-type and regional structure",
    "Implications for monitoring, carbon assessment, and risk",
    "Limitations",
    "Conclusion",
    "References",
}


def text_from_p(p_elm) -> str:
    return "".join(t.text or "" for t in p_elm.iter() if t.tag.endswith("}t"))


def set_run(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic


def format_paragraph(p, is_heading=False, is_title=False, italic=False):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_run(run, size=14 if is_title else 12, bold=is_heading or is_title, italic=italic)


def add_text_paragraph(doc: Document, text: str, is_title=False):
    p = doc.add_paragraph()
    p.add_run(text)
    stripped = text.strip()
    is_heading = stripped in HEADING_TEXT
    italic = stripped.startswith("Figure ") or stripped.startswith("Table ")
    format_paragraph(p, is_heading=is_heading, is_title=is_title, italic=italic)
    return p


def set_cell_text(cell, text: str, header=False):
    cell.text = text
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        for run in p.runs:
            set_run(run, size=9 if header else 8, bold=header)


def copy_table(dst: Document, src_table):
    table = dst.add_table(rows=1, cols=len(src_table.columns))
    table.style = "Table Grid"
    for j, cell in enumerate(src_table.rows[0].cells):
        set_cell_text(table.rows[0].cells[j], cell.text, header=True)
    for row in src_table.rows[1:]:
        cells = table.add_row().cells
        for j, cell in enumerate(row.cells):
            set_cell_text(cells[j], cell.text, header=False)
    dst.add_paragraph()


def setup_document(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Times New Roman"
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.size = Pt(12)


def rebuild() -> None:
    src = Document(SRC)
    dst = Document()
    setup_document(dst)
    body = src.element.body
    table_idx = 0
    title_written = False
    for child in body:
        if child.tag.endswith("}p"):
            text = text_from_p(child).strip()
            if not text:
                continue
            if text.startswith("Figure "):
                for prefix, img in FIGURE_IMAGES.items():
                    if text.startswith(prefix) and img.exists():
                        pic = dst.add_picture(str(img), width=Inches(6.6))
                        dst.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        break
            add_text_paragraph(dst, text, is_title=not title_written)
            title_written = True
        elif child.tag.endswith("}tbl"):
            copy_table(dst, src.tables[table_idx])
            table_idx += 1
    dst.save(OUT)


if __name__ == "__main__":
    rebuild()
