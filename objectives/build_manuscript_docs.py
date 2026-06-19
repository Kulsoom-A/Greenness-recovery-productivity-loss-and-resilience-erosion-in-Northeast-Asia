from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
OBJ1 = ROOT / "Paper2_Objectives" / "01_legacy_effects_kNDVI_GPP"
OBJ2 = ROOT / "Paper2_Objectives" / "02_event_order_resistance_recovery"
OBJ3 = ROOT / "Paper2_Objectives" / "03_vulnerability_hotspots_resilience_erosion"
OBJ4 = ROOT / "Paper2_Objectives" / "04_kNDVI_GPP_recovery_mismatch"
OBJ5 = ROOT / "Paper2_Objectives" / "05_drivers_controls_legacy_loss"
OBJ6 = ROOT / "Paper2_Objectives" / "06_sif_productivity_validation"
OBJ99 = ROOT / "Paper2_Objectives" / "99_validation_robustness"
OUT = ROOT / "Manuscript"
PUBFIG = OUT / "figures_publication"


def set_run(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph(doc, text="", style=None, size=11, bold=False, italic=False, align=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, italic=italic)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.08
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text)
    color = (31, 78, 121) if level == 1 else (68, 68, 68)
    set_run(r, size=15 if level == 1 else 12.5, bold=True, color=color)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(5)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size=9.5, bold=False, italic=True, color=(80, 80, 80))
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(9)
    return p


def add_qgis_placeholder(doc, label):
    p = doc.add_paragraph()
    r = p.add_run(f"[Insert QGIS-prepared spatial panel: {label}]")
    set_run(r, size=10, italic=True, color=(120, 120, 120))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    return p


def add_figure_placeholder(doc, label):
    p = doc.add_paragraph()
    r = p.add_run(f"[Insert final {label} here]")
    set_run(r, size=10, italic=True, color=(120, 120, 120))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(18)
    return p


def add_dataset_table(doc):
    rows = [
        (
            "kNDVI",
            "Annual vegetation greenness index",
            "2001-2024",
            "0.1 degree analysis grid",
            "Reference grid; primary greenness response variable for anomalies and recovery metrics; masked to vegetated study-area pixels.",
        ),
        (
            "MOD17 GPP",
            "Annual gross primary productivity",
            "2001-2024",
            "MODIS annual product, aligned to 0.1 degree",
            "Primary productivity response variable for cumulative loss, legacy duration, and driver models; bilinearly aligned before anomaly calculation.",
        ),
        (
            "GOSIF",
            "Annual solar-induced fluorescence",
            "2001-2024",
            "0.05 degree native, aggregated to 0.1 degree",
            "Independent productivity proxy for testing whether kNDVI-GPP mismatch depends on MOD17; aggregated and masked to the vegetation domain.",
        ),
        (
            "MOD13A2 EVI and NIRv",
            "Alternative greenness and canopy-structure proxies",
            "2001-2024",
            "MODIS, aggregated to 0.1 degree",
            "Structural greenness-side validation; QA-screened, with NIRv calculated as NDVI multiplied by near-infrared reflectance.",
        ),
        (
            "SPEI-12",
            "Monthly standardized precipitation evapotranspiration index",
            "2001-2022",
            "0.1 degree analysis grid",
            "Cumulative drought-stress criterion and event-window water-stress driver; December SPEI-12 used for annual drought status.",
        ),
        (
            "TerraClimate",
            "Tmax, precipitation, PET, climatic water deficit, and PDSI",
            "2001-2024",
            "Annual layers aligned to 0.1 degree",
            "Heat-stress threshold and climatic drivers; Tmax converted to pixel-wise z scores, other variables summarized across event years.",
        ),
        (
            "LC5 vegetation mask",
            "ENT, EBT, DNT, DBT, SHB, and GRS classes",
            "2024",
            "0.1 degree",
            "Vegetation-domain restriction and class-stratified summaries; non-vegetated and invalid pixels excluded.",
        ),
        (
            "MODIS MCD64A1 burned area",
            "Burned-area disturbance screen",
            "2004-2024",
            "Aligned to 0.1 degree",
            "External disturbance control; burned pixels separated from undisturbed-screen event pixels.",
        ),
        (
            "Hansen Global Forest Change",
            "Tree-cover loss and tree-cover baseline",
            "2001-2024",
            "30 m native, screened on 0.1 degree grid",
            "Forest-loss disturbance control and external hotspot-validation check; used as a conservative coarse screen for any loss occurrence.",
        ),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Dataset", "Variable", "Period", "Resolution", "Analytical use and preprocessing"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=7.2, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run(run, size=6.8)
    doc.add_paragraph()


def add_key_table(doc, df):
    cols = [
        "vegetation_class",
        "n_pixels_with_events",
        "event_frequency_mean",
        "kNDVI_legacy_years_t1_t4_mean",
        "GPP_legacy_years_t1_t4_mean",
    ]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    headers = [
        "Class",
        "Pixels with events",
        "Mean event frequency",
        "Mean kNDVI legacy years",
        "Mean GPP legacy years",
    ]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.5, bold=True)
    for _, row in df[cols].iterrows():
        cells = table.add_row().cells
        values = [
            str(row["vegetation_class"]),
            f"{int(row['n_pixels_with_events']):,}",
            f"{float(row['event_frequency_mean']):.2f}",
            f"{float(row['kNDVI_legacy_years_t1_t4_mean']):.2f}",
            f"{float(row['GPP_legacy_years_t1_t4_mean']):.2f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=8.5)
    doc.add_paragraph()


def add_event_order_table(doc, df):
    cols = [
        "vegetation_class",
        "n_recurrent_pixels",
        "event_count_mean",
        "kNDVI_delta_recovery_time_mean",
        "kNDVI_delta_cumulative_loss_mean",
        "GPP_delta_recovery_time_mean",
        "GPP_delta_cumulative_loss_mean",
    ]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    headers = [
        "Class",
        "Recurrent pixels",
        "Mean events",
        "Delta kNDVI recovery",
        "Delta kNDVI loss",
        "Delta GPP recovery",
        "Delta GPP loss",
    ]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=7.6, bold=True)
    for _, row in df[cols].iterrows():
        cells = table.add_row().cells
        values = [
            str(row["vegetation_class"]),
            f"{int(row['n_recurrent_pixels']):,}",
            f"{float(row['event_count_mean']):.2f}",
            f"{float(row['kNDVI_delta_recovery_time_mean']):.2f}",
            f"{float(row['kNDVI_delta_cumulative_loss_mean']):.3f}",
            f"{float(row['GPP_delta_recovery_time_mean']):.2f}",
            f"{float(row['GPP_delta_cumulative_loss_mean']):.1f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=7.4)
    doc.add_paragraph()


def add_uncertainty_table(doc, df):
    keep = df[df["vegetation_class"].astype(str) == "all_recurrent_forest"].copy()
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Metric", "n", "Mean", "95% CI low", "95% CI high"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.3, bold=True)
    for _, row in keep.iterrows():
        cells = table.add_row().cells
        values = [
            str(row["metric"]),
            f"{int(row['n']):,}",
            f"{float(row['mean']):.3f}" if abs(float(row["mean"])) < 10 else f"{float(row['mean']):.1f}",
            f"{float(row['ci_low']):.3f}" if abs(float(row["ci_low"])) < 10 else f"{float(row['ci_low']):.1f}",
            f"{float(row['ci_high']):.3f}" if abs(float(row["ci_high"])) < 10 else f"{float(row['ci_high']):.1f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=8)
    doc.add_paragraph()


def add_spatial_block_bootstrap_table(doc, df, metric_names=None):
    keep = df.copy()
    if metric_names is not None:
        keep = keep[keep["metric"].isin(metric_names)].copy()
    labels = {
        "kNDVI_later_minus_first_recovery_time": "kNDVI later-minus-first recovery time",
        "kNDVI_later_minus_first_cumulative_loss": "kNDVI later-minus-first cumulative loss",
        "GPP_later_minus_first_recovery_time": "GPP later-minus-first recovery time",
        "GPP_later_minus_first_cumulative_loss": "GPP later-minus-first cumulative loss",
        "kNDVI_cumulative_loss": "kNDVI cumulative loss",
        "GPP_cumulative_loss": "GPP cumulative loss",
        "kNDVI_legacy_duration": "kNDVI legacy duration",
        "GPP_legacy_duration": "GPP legacy duration",
        "hidden_mismatch_duration": "Hidden mismatch duration",
        "hidden_mismatch_prevalence": "Hidden mismatch prevalence",
    }
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Metric", "Pixels", "Blocks", "Mean", "95% CI low", "95% CI high"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.0, bold=True)
    for _, row in keep.iterrows():
        cells = table.add_row().cells
        mean = float(row["observed_mean"])
        low = float(row["ci_low"])
        high = float(row["ci_high"])
        values = [
            labels.get(str(row["metric"]), str(row["metric"])),
            f"{int(row['n_pixels']):,}",
            f"{int(row['n_blocks']):,}",
            f"{mean:.3f}" if abs(mean) < 10 else f"{mean:.1f}",
            f"{low:.3f}" if abs(low) < 10 else f"{low:.1f}",
            f"{high:.3f}" if abs(high) < 10 else f"{high:.1f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=7.8)
    doc.add_paragraph()


def add_clean_gap_table(doc, df):
    cols = [
        "vegetation_class",
        "n_clean_gap_recurrent_pixels",
        "event_count_mean",
        "GPP_first_cumulative_loss_mean",
        "GPP_later_cumulative_loss_mean",
        "GPP_delta_recovery_time_mean",
        "GPP_delta_cumulative_loss_mean",
    ]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    headers = [
        "Class",
        "Clean-gap pixels",
        "Mean events",
        "First GPP loss",
        "Later GPP loss",
        "Delta GPP recovery",
        "Delta GPP loss",
    ]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=7.6, bold=True)
    for _, row in df[cols].iterrows():
        cells = table.add_row().cells
        values = [
            str(row["vegetation_class"]),
            f"{int(row['n_clean_gap_recurrent_pixels']):,}",
            f"{float(row['event_count_mean']):.2f}",
            f"{float(row['GPP_first_cumulative_loss_mean']):.1f}",
            f"{float(row['GPP_later_cumulative_loss_mean']):.1f}",
            f"{float(row['GPP_delta_recovery_time_mean']):.2f}",
            f"{float(row['GPP_delta_cumulative_loss_mean']):.1f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=7.2)
    doc.add_paragraph()


def add_simple_dataframe_table(doc, df, columns, headers=None, percent_cols=None, max_rows=None):
    work = df.copy()
    if max_rows is not None:
        work = work.head(max_rows)
    headers = headers or columns
    percent_cols = set(percent_cols or [])
    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = str(text)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=7.6, bold=True)
    for _, row in work.iterrows():
        cells = table.add_row().cells
        values = []
        for col in columns:
            val = row[col]
            if pd.isna(val):
                values.append("")
            elif col in percent_cols:
                values.append(f"{float(val) * 100:.1f}%")
            elif isinstance(val, (int,)) or (hasattr(val, "is_integer") and float(val).is_integer() and "mean" not in col and "prevalence" not in col and "importance" not in col and "r2" not in col):
                values.append(f"{int(val):,}")
            elif isinstance(val, (float,)):
                values.append(f"{float(val):.3f}" if abs(float(val)) < 10 else f"{float(val):.1f}")
            else:
                values.append(str(val))
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=7.2)
    doc.add_paragraph()


def add_mismatch_ci_table(doc, df):
    keep = df[df["vegetation_class"].astype(str) == "all_event_forest"].copy()
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Metric", "n", "Mean", "95% CI low", "95% CI high"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.3, bold=True)
    for _, row in keep.iterrows():
        cells = table.add_row().cells
        values = [
            str(row["metric"]),
            f"{int(row['n']):,}",
            f"{float(row['mean']):.3f}",
            f"{float(row['ci_low']):.3f}",
            f"{float(row['ci_high']):.3f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=8)
    doc.add_paragraph()


def add_mismatch_table(doc, df):
    cols = [
        "vegetation_class",
        "n_pixels",
        "hidden_mismatch_prevalence",
        "hidden_mismatch_duration_mean",
        "hidden_fraction_of_GPP_suppression_mean",
    ]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    headers = ["Class", "Pixels", "Mismatch prevalence", "Mismatch duration", "Hidden fraction"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.2, bold=True)
    for _, row in df[cols].iterrows():
        cells = table.add_row().cells
        values = [
            str(row["vegetation_class"]),
            f"{int(row['n_pixels']):,}",
            f"{float(row['hidden_mismatch_prevalence']):.3f}",
            f"{float(row['hidden_mismatch_duration_mean']):.2f}",
            f"{float(row['hidden_fraction_of_GPP_suppression_mean']):.3f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=8)
    doc.add_paragraph()


def add_hotspot_summary_table(doc, df):
    keep = df[df["vegetation_class"].astype(str) == "all_event_forest"].copy()
    cols = ["hotspot_class", "n_pixels", "area_km2", "area_percent", "vulnerability_index_mean"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    headers = ["Hotspot class", "Pixels", "Area (km2)", "Area share (%)", "Mean vulnerability"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.2, bold=True)
    for _, row in keep[cols].iterrows():
        cells = table.add_row().cells
        values = [
            str(row["hotspot_class"]),
            f"{int(row['n_pixels']):,}",
            f"{float(row['area_km2']):,.0f}",
            f"{float(row['area_percent']):.1f}",
            f"{float(row['vulnerability_index_mean']):.3f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=8)
    doc.add_paragraph()


def add_vulnerability_table(doc, df):
    cols = ["vegetation_class", "n_pixels", "area_km2", "vulnerability_index_mean", "severe_hotspot_percent"]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    headers = ["Class", "Pixels", "Area (km2)", "Mean vulnerability", "Severe hotspots (%)"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.2, bold=True)
    for _, row in df[cols].iterrows():
        cells = table.add_row().cells
        values = [
            str(row["vegetation_class"]),
            f"{int(row['n_pixels']):,}",
            f"{float(row['area_km2']):,.0f}",
            f"{float(row['vulnerability_index_mean']):.3f}",
            f"{float(row['severe_hotspot_percent']):.1f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=8)
    doc.add_paragraph()


def add_model_performance_table(doc, df):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Outcome", "Model", "Test n", "R2/AUC", "MAE"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.2, bold=True)
    for _, row in df.iterrows():
        metric = row["r2"] if pd.notna(row.get("r2")) else row.get("auc")
        cells = table.add_row().cells
        values = [
            str(row["outcome"]),
            str(row["model"]).replace("RandomForest", "RF "),
            f"{int(row['n_test']):,}",
            f"{float(metric):.3f}",
            "" if pd.isna(row.get("mae")) else f"{float(row['mae']):.3f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=8)
    doc.add_paragraph()


def add_model_validation_table(doc, random_df, spatial_df):
    rows = []
    for _, row in random_df.iterrows():
        metric = row["r2"] if pd.notna(row.get("r2")) else row.get("auc")
        rows.append(
            {
                "validation": "Random held-out",
                "outcome": row["outcome"],
                "test": int(row["n_test"]),
                "metric": metric,
                "mae": row.get("mae"),
                "folds": "",
            }
        )
    for _, row in spatial_df.iterrows():
        metric = row["r2_mean"] if pd.notna(row.get("r2_mean")) else row.get("auc_mean")
        rows.append(
            {
                "validation": "Spatial block",
                "outcome": row["outcome"],
                "test": int(round(float(row["mean_n_test"]))),
                "metric": metric,
                "mae": row.get("mae_mean"),
                "folds": int(row["n_folds"]),
            }
        )
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Validation", "Outcome", "Test n", "Folds", "R2/AUC", "MAE"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.0, bold=True)
    for row in rows:
        cells = table.add_row().cells
        values = [
            str(row["validation"]),
            str(row["outcome"]),
            f"{int(row['test']):,}",
            "" if row["folds"] == "" else str(row["folds"]),
            f"{float(row['metric']):.3f}",
            "" if pd.isna(row["mae"]) else f"{float(row['mae']):.3f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=7.5)
    doc.add_paragraph()


def add_spatial_cv_fold_table(doc, df):
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["Fold", "Outcome", "Test n", "Blocks", "R2", "MAE", "AUC"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.0, bold=True)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        values = [
            str(int(row["fold"])),
            str(row["outcome"]),
            f"{int(row['n_test']):,}",
            str(int(row["test_blocks"])),
            "" if pd.isna(row.get("r2")) else f"{float(row['r2']):.3f}",
            "" if pd.isna(row.get("mae")) else f"{float(row['mae']):.3f}",
            "" if pd.isna(row.get("auc")) else f"{float(row['auc']):.3f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=7.2)
    doc.add_paragraph()


def add_importance_table(doc, gpp_df, severe_df):
    rows = []
    for label, df in [("GPP cumulative loss", gpp_df), ("Severe hotspot", severe_df)]:
        for _, row in df.head(5).iterrows():
            feature = str(row["feature"]).replace("vegetation_class_", "Veg: ")
            rows.append(
                {
                    "outcome": label,
                    "feature": feature,
                    "importance_mean": float(row["importance_mean"]),
                    "importance_sd": float(row["importance_sd"]),
                }
            )
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Outcome", "Driver", "Importance", "SD"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.2, bold=True)
    for row in rows:
        cells = table.add_row().cells
        values = [row["outcome"], row["feature"], f"{row['importance_mean']:.3f}", f"{row['importance_sd']:.3f}"]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=8)
    doc.add_paragraph()


def add_robustness_table(doc, greenness_df, loo_df, tree_df):
    kndvi = greenness_df[
        (greenness_df["greenness_index"].astype(str) == "kNDVI")
        & (greenness_df["vegetation_class"].astype(str) == "all_event_forest")
    ].iloc[0]
    ndvi = greenness_df[
        (greenness_df["greenness_index"].astype(str) == "NDVI")
        & (greenness_df["vegetation_class"].astype(str) == "all_event_forest")
    ].iloc[0]
    tree = tree_df[tree_df["domain"].astype(str) == "tree_classes_only"].iloc[0]
    all_domain = tree_df[tree_df["domain"].astype(str) == "all_event_vegetation"].iloc[0]
    rows = [
        ("kNDVI-GPP mismatch prevalence", f"{float(kndvi['hidden_mismatch_prevalence']):.3f}", "All event-affected pixels"),
        ("NDVI-GPP mismatch prevalence", f"{float(ndvi['hidden_mismatch_prevalence']):.3f}", "Independent greenness check"),
        ("Minimum leave-one-out hotspot Jaccard", f"{float(loo_df['jaccard_with_original_severe'].min()):.3f}", "Across omitted hotspot components"),
        ("Mean leave-one-out hotspot Jaccard", f"{float(loo_df['jaccard_with_original_severe'].mean()):.3f}", "Across omitted hotspot components"),
        ("Tree-only severe hotspot share", f"{float(tree['severe_hotspot_percent']):.1f}%", "ENT, EBT, DNT, DBT only"),
        ("All-domain severe hotspot share", f"{float(all_domain['severe_hotspot_percent']):.1f}%", "ENT, EBT, DNT, DBT, SHB, GRS"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Robustness check", "Value", "Interpretation domain"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.2, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=8)
    doc.add_paragraph()


def add_external_disturbance_table(doc, df):
    cols = [
        "domain",
        "n_pixels",
        "GPP_cumulative_loss_mean",
        "hidden_mismatch_prevalence",
        "severe_hotspot_percent",
        "vulnerability_index_mean",
    ]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    headers = ["Domain", "Pixels", "GPP loss", "Mismatch prevalence", "Severe hotspots", "Vulnerability"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=7.8, bold=True)
    for _, row in df[cols].iterrows():
        cells = table.add_row().cells
        values = [
            str(row["domain"]).replace("_", " "),
            f"{int(row['n_pixels']):,}",
            f"{float(row['GPP_cumulative_loss_mean']):.1f}",
            f"{float(row['hidden_mismatch_prevalence']):.3f}",
            f"{float(row['severe_hotspot_percent']):.1f}%",
            f"{float(row['vulnerability_index_mean']):.3f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=7.4)
    doc.add_paragraph()


def add_sif_validation_table(doc, df):
    cols = [
        "vegetation_class",
        "n_pixels",
        "GPP_hidden_mismatch_prevalence",
        "GOSIF_hidden_mismatch_prevalence",
        "both_GPP_and_GOSIF_mismatch",
        "duration_correlation",
    ]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    headers = ["Class", "Pixels", "MOD17 GPP mismatch", "GOSIF mismatch", "Both proxies", "Duration r"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=7.8, bold=True)
    for _, row in df[cols].iterrows():
        cells = table.add_row().cells
        values = [
            str(row["vegetation_class"]),
            f"{int(row['n_pixels']):,}",
            f"{float(row['GPP_hidden_mismatch_prevalence']):.3f}",
            f"{float(row['GOSIF_hidden_mismatch_prevalence']):.3f}",
            f"{float(row['both_GPP_and_GOSIF_mismatch']):.3f}",
            f"{float(row['duration_correlation']):.3f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=7.3)
    doc.add_paragraph()


def add_structural_greenness_table(doc, df):
    main = df[df["vegetation_class"].astype(str) == "all_event_forest"].copy()
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Greenness proxy", "Productivity proxy", "Pixels", "Mismatch prevalence", "Mismatch duration"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.0, bold=True)
    for _, row in main.iterrows():
        cells = table.add_row().cells
        values = [
            str(row["greenness_proxy"]),
            str(row["productivity_proxy"]),
            f"{int(row['n_pixels']):,}",
            f"{float(row['hidden_mismatch_prevalence']):.3f}",
            f"{float(row['hidden_mismatch_duration_mean']):.2f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=7.6)
    doc.add_paragraph()


def add_event_timing_null_table(doc, df):
    main = df[df["vegetation_class"].astype(str) == "all_event_forest"].copy()
    keep = [
        "kNDVI_cumulative_loss",
        "GPP_cumulative_loss",
        "kNDVI_legacy_duration",
        "GPP_legacy_duration",
        "kNDVI_GPP_hidden_mismatch_prevalence",
    ]
    labels = {
        "kNDVI_cumulative_loss": "kNDVI cumulative loss",
        "GPP_cumulative_loss": "GPP cumulative loss",
        "kNDVI_legacy_duration": "kNDVI legacy duration",
        "GPP_legacy_duration": "GPP legacy duration",
        "kNDVI_GPP_hidden_mismatch_prevalence": "Hidden mismatch prevalence",
    }
    main = main[main["metric"].isin(keep)].copy()
    main["metric_label"] = main["metric"].map(labels)
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Metric", "Observed", "Null mean", "Null 95% low", "Null 95% high", "Obs/null"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=7.6, bold=True)
    for _, row in main.iterrows():
        cells = table.add_row().cells
        values = [
            str(row["metric_label"]),
            f"{float(row['observed_mean']):.3f}",
            f"{float(row['null_mean']):.3f}",
            f"{float(row['null_ci_low']):.3f}",
            f"{float(row['null_ci_high']):.3f}",
            f"{float(row['observed_to_null_ratio']):.2f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=7.1)
    doc.add_paragraph()


def add_sensitivity_table(doc, df):
    cols = [
        "spei_threshold",
        "tmax_z_threshold",
        "n_recurrent_pixels",
        "kNDVI_delta_cumulative_loss_mean",
        "GPP_delta_cumulative_loss_mean",
    ]
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    headers = ["SPEI threshold", "Tmax z", "Recurrent pixels", "Delta kNDVI loss", "Delta GPP loss"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.2, bold=True)
    for _, row in df[cols].iterrows():
        cells = table.add_row().cells
        values = [
            f"{float(row['spei_threshold']):.1f}",
            f"{float(row['tmax_z_threshold']):.1f}",
            f"{int(row['n_recurrent_pixels']):,}",
            f"{float(row['kNDVI_delta_cumulative_loss_mean']):.3f}",
            f"{float(row['GPP_delta_cumulative_loss_mean']):.1f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=8)
    doc.add_paragraph()


def save_doc(doc, path: Path) -> Path:
    try:
        doc.save(path)
        return path
    except PermissionError:
        updated = path.with_name(path.stem + "_Updated" + path.suffix)
        doc.save(updated)
        return updated


def base_doc(title, subtitle):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    p = doc.add_paragraph()
    r = p.add_run(title)
    set_run(r, size=20, bold=True, color=(31, 78, 121))
    p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    r = p.add_run(subtitle)
    set_run(r, size=11, italic=True, color=(90, 90, 90))
    p.paragraph_format.space_after = Pt(12)
    return doc


def build_main():
    OUT.mkdir(parents=True, exist_ok=True)
    legacy = pd.read_csv(OBJ1 / "outputs" / "tables" / "objective1_legacy_summary_by_class.csv")
    order = pd.read_csv(OBJ2 / "outputs" / "tables" / "event_order_summary_by_class.csv")
    ci = pd.read_csv(OBJ2 / "outputs" / "tables" / "event_order_bootstrap_ci.csv")
    mismatch = pd.read_csv(OBJ4 / "outputs" / "tables" / "mismatch_summary_by_class.csv")
    mismatch_ci = pd.read_csv(OBJ4 / "outputs" / "tables" / "mismatch_bootstrap_ci.csv")
    hotspot = pd.read_csv(OBJ3 / "outputs" / "tables" / "hotspot_class_summary_by_vegetation.csv")
    vulnerability = pd.read_csv(OBJ3 / "outputs" / "tables" / "vulnerability_index_summary_by_vegetation.csv")
    model_perf = pd.read_csv(OBJ5 / "outputs" / "tables" / "model_performance_summary.csv")
    spatial_cv = pd.read_csv(OBJ5 / "outputs" / "tables" / "spatial_block_cv_summary.csv")
    gpp_importance = pd.read_csv(OBJ5 / "outputs" / "tables" / "rf_importance_GPP_cumulative_loss.csv")
    severe_importance = pd.read_csv(OBJ5 / "outputs" / "tables" / "rf_importance_severe_hotspot.csv")
    sif_comparison = pd.read_csv(OBJ6 / "outputs" / "tables" / "gpp_vs_gosif_mismatch_comparison.csv")
    sif_summary = pd.read_csv(OBJ6 / "outputs" / "tables" / "gosif_mismatch_summary_by_class.csv")
    structural_greenness = pd.read_csv(OBJ6 / "outputs" / "tables" / "structural_greenness_gosif_validation.csv")
    robust_greenness = pd.read_csv(OBJ99 / "outputs" / "tables" / "greenness_index_mismatch_robustness.csv")
    robust_loo = pd.read_csv(OBJ99 / "outputs" / "tables" / "hotspot_leave_one_component_out_stability.csv")
    robust_tree = pd.read_csv(OBJ99 / "outputs" / "tables" / "tree_only_domain_robustness.csv")
    null_summary = pd.read_csv(OBJ99 / "outputs" / "tables" / "event_timing_null_summary.csv")
    clean_gap = pd.read_csv(OBJ2 / "outputs" / "tables" / "clean_gap_event_order_summary.csv")
    spatial_desc = pd.read_csv(OBJ99 / "outputs" / "tables" / "spatial_block_descriptive_bootstrap.csv")
    fractional_imp = pd.read_csv(OBJ99 / "outputs" / "tables" / "rf_importance_GPP_fractional_cumulative_loss.csv")
    fractional_perf = pd.read_csv(OBJ99 / "outputs" / "tables" / "fractional_loss_model_performance.csv")
    event_type_response = pd.read_csv(OBJ99 / "outputs" / "tables" / "compound_vs_univariate_event_response.csv")
    hotspot_hansen_validation = pd.read_csv(OBJ99 / "outputs" / "tables" / "hotspot_hansen_forest_loss_validation.csv")
    disturbance = pd.read_csv(OBJ99 / "outputs" / "tables" / "external_disturbance_control_summary.csv")
    all_forest = legacy[legacy["vegetation_class"].astype(str) == "all_forest"].iloc[0]
    recurrent = order[order["vegetation_class"].astype(str) == "all_recurrent_forest"].iloc[0]
    gpp_ci = spatial_desc[spatial_desc["metric"].astype(str) == "GPP_later_minus_first_cumulative_loss"].iloc[0]
    clean_gap_all = clean_gap[clean_gap["vegetation_class"].astype(str) == "all_clean_gap_recurrent_forest"].iloc[0]
    mismatch_all = mismatch[mismatch["vegetation_class"].astype(str) == "all_event_forest"].iloc[0]
    severe_hotspot = hotspot[
        (hotspot["vegetation_class"].astype(str) == "all_event_forest")
        & (hotspot["hotspot_class"].astype(str) == "Severe multi-metric hotspot")
    ].iloc[0]
    vulnerability_all = vulnerability[vulnerability["vegetation_class"].astype(str) == "all_event_forest"].iloc[0]
    dbt_vulnerability = vulnerability[vulnerability["vegetation_class"].astype(str) == "DBT"].iloc[0]
    dbt_mismatch = mismatch[mismatch["vegetation_class"].astype(str) == "DBT"].iloc[0]
    gpp_perf = model_perf[model_perf["outcome"].astype(str) == "GPP_cumulative_loss"].iloc[0]
    hotspot_perf = model_perf[model_perf["outcome"].astype(str) == "severe_hotspot"].iloc[0]
    gpp_spatial = spatial_cv[spatial_cv["outcome"].astype(str) == "GPP_cumulative_loss"].iloc[0]
    hotspot_spatial = spatial_cv[spatial_cv["outcome"].astype(str) == "severe_hotspot"].iloc[0]
    sif_all = sif_comparison[sif_comparison["vegetation_class"].astype(str) == "all_event_forest"].iloc[0]
    sif_summary_all = sif_summary[sif_summary["vegetation_class"].astype(str) == "all_event_forest"].iloc[0]
    structural_all = structural_greenness[structural_greenness["vegetation_class"].astype(str) == "all_event_forest"]
    evi_structural = structural_all[structural_all["greenness_proxy"].astype(str) == "EVI"].iloc[0]
    nirv_structural = structural_all[structural_all["greenness_proxy"].astype(str) == "NIRv"].iloc[0]
    top_gpp_driver = gpp_importance.iloc[0]
    top_hotspot_driver = severe_importance.iloc[0]
    top_fractional_driver = fractional_imp.iloc[0]
    fractional_baseline = fractional_imp[fractional_imp["feature"].astype(str) == "pre_event_gpp_baseline"].iloc[0]
    fractional_spatial = fractional_perf[fractional_perf["model"].astype(str) == "RandomForestRegressor_spatial_block"].iloc[0]
    compound_event = event_type_response[event_type_response["event_type"].astype(str) == "compound"].iloc[0]
    drought_only_event = event_type_response[event_type_response["event_type"].astype(str) == "drought_only"].iloc[0]
    heat_only_event = event_type_response[event_type_response["event_type"].astype(str) == "heat_only"].iloc[0]
    hansen_severe = hotspot_hansen_validation[hotspot_hansen_validation["domain"].astype(str) == "severe_hotspot"].iloc[0]
    hansen_nonsevere = hotspot_hansen_validation[hotspot_hansen_validation["domain"].astype(str) == "non_severe_event_pixels"].iloc[0]
    hansen_model = hotspot_hansen_validation[hotspot_hansen_validation["domain"].astype(str) == "predictive_model_summary"].iloc[0]
    ndvi_robust = robust_greenness[
        (robust_greenness["greenness_index"].astype(str) == "NDVI")
        & (robust_greenness["vegetation_class"].astype(str) == "all_event_forest")
    ].iloc[0]
    tree_robust = robust_tree[robust_tree["domain"].astype(str) == "tree_classes_only"].iloc[0]
    null_all = null_summary[null_summary["vegetation_class"].astype(str) == "all_event_forest"]
    null_gpp_loss = null_all[null_all["metric"].astype(str) == "GPP_cumulative_loss"].iloc[0]
    null_gpp_duration = null_all[null_all["metric"].astype(str) == "GPP_legacy_duration"].iloc[0]
    null_k_duration = null_all[null_all["metric"].astype(str) == "kNDVI_legacy_duration"].iloc[0]
    null_mismatch_prev = null_all[null_all["metric"].astype(str) == "kNDVI_GPP_hidden_mismatch_prevalence"].iloc[0]
    undisturbed = disturbance[disturbance["domain"].astype(str) == "undisturbed_screen"].iloc[0]
    mismatch_prev_ci = spatial_desc[spatial_desc["metric"].astype(str) == "hidden_mismatch_prevalence"].iloc[0]
    fig = PUBFIG / "figure1_event_panel_ggstyle.png"
    spatial_fig2 = PUBFIG / "qgis_guide_figure2_recurrent_resilience_erosion.png"
    spatial_fig3 = PUBFIG / "qgis_guide_figure3_hidden_productivity_mismatch.png"
    spatial_fig4 = PUBFIG / "qgis_guide_figure4_integrated_vulnerability_hotspots.png"
    spatial_fig5 = PUBFIG / "qgis_guide_figure5_spatial_controls_predicted_vulnerability.png"
    driver_fig = PUBFIG / "figure11_driver_importance_combined.png"
    robustness_fig = PUBFIG / "figure12_internal_robustness_checks.png"
    disturbance_fig = PUBFIG / "figure13_external_disturbance_control.png"
    sif_fig = PUBFIG / "figure14_sif_productivity_validation.png"
    structural_fig = PUBFIG / "figure15_structural_greenness_sif_validation.png"
    null_fig = PUBFIG / "figure16_event_timing_null_model.png"

    doc = base_doc(
        "From Greening-Masked Vulnerability to Resilience Erosion",
        "Main manuscript working draft",
    )
    add_heading(doc, "Materials and methods", 1)
    add_heading(doc, "Study area", 2)
    add_paragraph(
        doc,
        "Northeast Asia was selected as the study domain because it contains one of the largest climatically transitional vegetation regions of the Northern Hemisphere, linking boreal and cold-temperate forests, montane forest belts, shrublands, and dry grassland margins within a single heat- and water-stress gradient. "
        "This geography is scientifically useful for testing compound hot-dry legacies because ecosystem productivity in the region is controlled by both growing-season thermal limitation and episodic water deficit, while recent warming has increased the likelihood that drought years also occur under anomalously high temperature. "
        "The domain also includes strong contrasts in canopy structure and plant functional type, allowing the analysis to test whether greenness recovery, productivity suppression, recurrent-event sensitivity, and vulnerability hotspots are consistent across evergreen needleleaf trees (ENT), evergreen broadleaf trees (EBT), deciduous needleleaf trees (DNT), deciduous broadleaf trees (DBT), shrubs (SHB), and grasslands (GRS). "
        "All analyses were clipped to the project boundary layer and harmonized to the common 0.1 degree grid shown in Figure 1, then restricted to vegetated pixels with complete greenness, productivity, climate, and disturbance-screen information.",
    )
    add_figure_placeholder(doc, "Figure 1 study-area map")
    add_caption(
        doc,
        "Figure 1. Study area and analysis domain in Northeast Asia. "
        "The final map should show the project boundary, vegetation classes used in the analysis, major geographic context, and the common 0.1 degree analysis grid or representative grid extent.",
    )
    add_heading(doc, "Datasets and preprocessing", 2)
    add_paragraph(
        doc,
        "Annual kNDVI, EVI, NIRv, monthly SPEI-12, TerraClimate maximum temperature, precipitation, potential evapotranspiration, climatic water deficit, PDSI, annual MOD17 GPP, annual GOSIF solar-induced fluorescence, and the six-class vegetation mask were harmonized to a common 0.1 degree EPSG:4326 grid. "
        "The kNDVI and climate products formed the reference grid. MOD17 GPP was bilinearly aligned to this grid before anomaly calculation, while categorical vegetation data were preserved as class codes. "
        "The annual GOSIF product was decompressed, averaged from its native 0.05 degree grid to the 0.1 degree analysis grid, and masked to the study boundary and vegetation domain. "
        "EVI was derived from MOD13A2, and NIRv was calculated as NDVI multiplied by near-infrared reflectance using QA-screened MOD13A2 red and near-infrared reflectance before aggregation to the analysis grid. "
        "Annual climate summaries were calculated from the TerraClimate layers, and December SPEI-12 was used to represent cumulative water stress at the end of each growing season and calendar year. "
        "All raster calculations were performed pixel by pixel so that event exposure, response baselines, legacy metrics, and driver variables were spatially matched. "
        "The datasets, periods, analytical roles, and preprocessing steps are summarized in Table 1.",
    )
    add_dataset_table(doc)
    add_caption(doc, "Table 1. Datasets and preprocessing steps used for event definition, ecosystem-response metrics, robustness tests, and disturbance screening.")
    add_heading(doc, "Compound hot-dry events and response metrics", 2)
    add_paragraph(
        doc,
        "Compound hot-dry events were identified independently for each pixel and year. "
        "A year was classified as an event when December SPEI-12 was <= -1.0 and annual Tmax exceeded the local 2001-2024 mean by at least +1 standard deviation. "
        "The heat threshold was calculated as Tmax_z(i,y) = [Tmax(i,y) - mean(Tmax_i)] / sd(Tmax_i), where i denotes pixel and y denotes year. "
        "To isolate the compound-event interpretation, two mutually exclusive univariate event sets were also evaluated with identical response windows: drought-only years satisfied SPEI-12 <= -1.0 but Tmax_z < +1, and heat-only years satisfied Tmax_z >= +1 but SPEI-12 > -1.0. "
        "Event years were limited to 2004-2020 so that every event had a complete pre-event baseline window from t-3 to t-1 and a post-event response window from t+1 to t+4. "
        "For a response variable Y, the pre-event baseline for pixel i and event year e was B(i,e) = mean[Y(i,e-3), Y(i,e-2), Y(i,e-1)]. "
        "Relative-year anomalies were then calculated as A(i,e,tau) = Y(i,e+tau) - B(i,e), where tau ranges from -3 to +4 and tau = 0 is the event year. "
        "Cumulative loss was defined as L(i,e) = -sum_tau min[A(i,e,tau), 0] for tau = 0...4, so larger positive values indicate greater cumulative suppression. "
        "Legacy duration was defined as D(i,e) = sum I[A(i,e,tau) < 0] for tau = 1...4. "
        "Pixel-level maps summarize event-frequency-weighted mean loss and duration across valid events. "
        "Because these one-sided metrics can be positive even under baseline-centered variability, a circular-shift event-timing null model was generated by randomly shifting each pixel's annual kNDVI and GPP time series by a nonzero offset while holding the observed event years fixed; observed loss, duration, and mismatch metrics were then compared against 100 shifted realizations.",
    )
    add_heading(doc, "Recurrent exposure, mismatch, and hotspot metrics", 2)
    add_paragraph(
        doc,
        "To test whether repeated exposure was associated with resilience erosion, events were ranked chronologically within each pixel. "
        "The first valid event was compared against the mean of all later valid events at the same pixel. "
        "Resistance was R(i,e) = A(i,e,0), recovery time was T(i,e) = min tau such that A(i,e,tau) >= 0 for tau = 1...4, and unrecovered events were assigned T = 5. "
        "Later-minus-first differences were calculated for resistance, recovery time, and cumulative loss; positive recovery-time and cumulative-loss differences indicate slower or stronger later-event impacts. "
        "To test whether later-event baselines were contaminated by earlier response windows, a clean-gap sensitivity analysis repeated the recurrent-event comparison using only later events separated by at least eight years from the previous valid event, preventing the t-3 to t-1 baseline from overlapping the previous event's t0 to t+4 response window. "
        "Hidden productivity-loss mismatch was designed to identify cases where greenness recovery would imply ecosystem recovery but productivity remained suppressed. "
        "For each post-event year tau = 1...4, a hidden mismatch occurred when A_kNDVI(i,e,tau) >= 0 and A_GPP(i,e,tau) < 0. "
        "Mismatch duration was H(i) = sum I[A_kNDVI >= 0 and A_GPP < 0] across post-event years and valid events, while prevalence was the fraction of event-affected pixels with H > 0. "
        "A standardized mismatch-intensity metric was also calculated as M(i,tau) = z[A_kNDVI(i,tau)] - z[A_GPP(i,tau)], averaged across post-event years. "
        "A vulnerability-hotspot classification was then derived by integrating exposure, productivity loss, legacy duration, recurrent-event change, and kNDVI-GPP mismatch metrics. "
        "Each continuous component was scaled from 0 to 1 using the 5th and 95th percentiles across event-affected forest pixels, and the composite vulnerability index was calculated as the mean of the scaled components. "
        "Formally, V(i) = mean[S_m(i)], where S_m is the percentile-scaled value of metric m after clipping to the 0-1 interval. "
        "The components included event frequency, GPP cumulative loss, GPP legacy duration, later-minus-first GPP cumulative-loss change, later-minus-first GPP recovery-time change, hidden mismatch duration, and standardized mismatch intensity. "
        "Pixels were assigned to ordered classes representing low-impact recovery, delayed productivity recovery, hidden productivity suppression, recurrent resilience erosion, and severe multi-metric hotspots.",
    )
    add_heading(doc, "Driver modelling and robustness checks", 2)
    add_paragraph(
        doc,
        "Potential climatic and ecological controls were evaluated using event-window driver layers. "
        "For each event-affected pixel, event-year means were calculated for SPEI-12, Tmax z-score, climatic water-deficit z-score, precipitation z-score, PET z-score, PDSI, event frequency, and pre-event GPP baseline. "
        "Associations were screened using Spearman rank correlations, while nonlinear multivariate controls were estimated with random-forest regression for GPP cumulative loss and vulnerability index, and random-forest classification for severe hotspot occurrence. "
        "Models were fit on a fixed random sample of event-affected pixels, evaluated on held-out test data, and interpreted using permutation importance. "
        "To test whether model skill was inflated by spatial autocorrelation, the same predictor set was also evaluated with five-fold spatial block cross-validation using 10 degree geographic blocks as held-out groups. "
        "Descriptive uncertainty for prevalence, legacy, mismatch, and recurrent-event statistics was recalculated using a 10 degree spatial-block bootstrap, in which geographic blocks rather than individual pixels were resampled with replacement. "
        "Because absolute cumulative GPP loss can mechanically scale with pre-event productivity, GPP cumulative loss was also normalized by five pre-event baseline years and the driver model was re-fit to this fractional-loss outcome. "
        "A logistic model with robust standard errors was used as a supplementary parametric check for severe-hotspot odds. "
        "Three internal robustness checks were used to assess whether the main conclusions depended on a single index, a single hotspot component, or the inclusion of non-tree vegetation classes. "
        "First, NDVI was substituted for kNDVI in the hidden productivity-loss mismatch definition while keeping the same GPP anomalies and event windows. "
        "Second, the severe-hotspot domain was recalculated after omitting each vulnerability-index component in turn; each leave-one-component-out severe map was compared with the original severe-hotspot domain using Jaccard similarity and pixel agreement. "
        "Third, hotspot statistics were recalculated for tree classes only (ENT, EBT, DNT, and DBT) and compared with the full event-affected vegetation domain. "
        "Finally, the hidden productivity-loss mismatch was recalculated using GOSIF as an independent productivity proxy to test whether the kNDVI-productivity mismatch persisted outside the MOD17 light-use-efficiency formulation; the greenness side of the mismatch was then retested using EVI and NIRv to avoid relying only on the deterministic NDVI-to-kNDVI transformation. "
    )
    add_heading(doc, "External disturbance control", 2)
    add_paragraph(
        doc,
        "External disturbance was screened using MODIS MCD64A1 burned area and Hansen Global Forest Change tree-cover loss products downloaded from Google Earth Engine and aligned to the 0.1 degree analysis grid. "
        "Burned pixels were identified for 2004-2024, covering the event and post-event response period. "
        "Forest-loss disturbance was identified from the Hansen annual loss-year layer over the same period. "
        "Pixels flagged by burned area or forest loss were separated from the remaining event-affected pixels, and key metrics were recalculated for all event pixels, disturbed pixels, and pixels passing the undisturbed-screen sensitivity test. "
        "The severe-hotspot domain was also compared with the Hansen forest-loss screen as an external validation target; this test was interpreted cautiously because the available 0.1 degree Hansen screen records whether any tree-cover loss occurred within a coarse grid cell rather than a fractional mortality response.",
    )

    add_heading(doc, "Results", 1)
    add_heading(doc, "Event legacies and recurrent exposure", 2)
    add_paragraph(
        doc,
        f"Across all forest classes, {int(all_forest['n_pixels_with_events']):,} pixels experienced at least one compound hot-dry event during 2004-2020. "
        f"The mean event frequency was {float(all_forest['event_frequency_mean']):.2f} events per affected pixel, with a maximum of {float(all_forest['event_frequency_max']):.0f} events. "
        f"The mean post-event legacy duration was {float(all_forest['kNDVI_legacy_years_t1_t4_mean']):.2f} years for kNDVI and {float(all_forest['GPP_legacy_years_t1_t4_mean']):.2f} years for GPP (Supplementary Table S26). "
        f"The circular-shift null model showed that the raw kNDVI duration was not above timing-randomized expectation (observed/null = {float(null_k_duration['observed_to_null_ratio']):.2f}), whereas GPP cumulative loss and GPP legacy duration exceeded the null expectation (loss observed/null = {float(null_gpp_loss['observed_to_null_ratio']):.2f}; duration observed/null = {float(null_gpp_duration['observed_to_null_ratio']):.2f}; Supplementary Fig. S13 and Supplementary Table S20). "
        f"The univariate-event comparison showed that compound events produced larger GPP cumulative loss than drought-only years ({float(compound_event['GPP_cumulative_loss_mean']):.1f} versus {float(drought_only_event['GPP_cumulative_loss_mean']):.1f} g C m-2), but heat-only years produced similarly strong or stronger losses ({float(heat_only_event['GPP_cumulative_loss_mean']):.1f} g C m-2; Supplementary Fig. S14 and Supplementary Table S24). "
        f"The same pattern appeared for fractional cumulative loss ({float(compound_event['GPP_fractional_cumulative_loss_mean']):.3f} for compound, {float(drought_only_event['GPP_fractional_cumulative_loss_mean']):.3f} for drought-only, and {float(heat_only_event['GPP_fractional_cumulative_loss_mean']):.3f} for heat-only), so the compound-event framing is interpreted as a hot-dry co-occurrence test rather than as evidence that compound years always exceed isolated heat years. "
        "Thus, absolute legacy interpretation is focused on productivity suppression and mismatch metrics rather than on the raw one-sided kNDVI duration.",
    )
    add_paragraph(
        doc,
        "The mean event-panel response (Figure 2) shows the temporal basis of this divergence. "
        "kNDVI anomalies returned close to baseline more rapidly, whereas GPP remained more persistently negative after the event year, especially during t+2 and t+3. "
        "The separation between the two trajectories therefore reflects a recovery mismatch rather than a short-lived event-year anomaly. "
        "The class-level legacy table and the supporting spatial map panel in Supplementary Fig. S1 show that this was not only an averaged time-series effect, but also a spatially structured legacy pattern across the study domain.",
    )
    doc.add_picture(str(fig), width=Inches(6.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(
        doc,
        "Figure 2. Mean event-panel anomalies for kNDVI and GPP across all affected forest pixels. "
        "Anomalies are relative to the pre-event baseline calculated from t-3 to t-1; the dashed vertical line marks the compound hot-dry event year.",
    )
    add_paragraph(
        doc,
        f"Among recurrent-event pixels, {int(recurrent['n_recurrent_pixels']):,} pixels experienced at least two compound hot-dry events, with a mean of {float(recurrent['event_count_mean']):.2f} events per recurrent pixel. "
        f"Later events produced slightly larger mean kNDVI cumulative loss than first events ({float(recurrent['kNDVI_later_cumulative_loss_mean']):.3f} versus {float(recurrent['kNDVI_first_cumulative_loss_mean']):.3f}). "
        f"The difference was much stronger for GPP, where mean cumulative loss increased from {float(recurrent['GPP_first_cumulative_loss_mean']):.1f} to {float(recurrent['GPP_later_cumulative_loss_mean']):.1f} g C m-2 accumulated across the response window. "
        f"The later-minus-first GPP cumulative-loss difference was {float(gpp_ci['observed_mean']):.1f}, with a spatial-block bootstrap 95% confidence interval of {float(gpp_ci['ci_low']):.1f} to {float(gpp_ci['ci_high']):.1f}. "
        f"This pattern was retained under the conservative clean-gap sensitivity test, where {int(clean_gap_all['n_clean_gap_recurrent_pixels']):,} recurrent pixels remained and later-event GPP cumulative loss exceeded first-event loss by {float(clean_gap_all['GPP_delta_cumulative_loss_mean']):.1f} g C m-2 (Supplementary Table S21). "
        "Thus, recurrent exposure did not simply reproduce the first-event response; it was associated with a larger productivity penalty during later events. "
        "These estimates are summarized in Supplementary Tables S2 and S22 and remained positive across the threshold-sensitivity matrix in Supplementary Table S4. "
        "Figure 3 emphasizes the overlap between event frequency, the recurrent-event domain, and later-minus-first GPP loss rather than presenting the bar charts as the primary evidence.",
    )
    doc.add_picture(str(spatial_fig2), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(
        doc,
        "Figure 3. Spatial structure of recurrent compound hot-dry exposure and resilience erosion. "
        "Panels show compound hot-dry event frequency (A), the recurrent-event domain (B), later-minus-first GPP cumulative loss (C), and later-minus-first GPP recovery time (D).",
    )
    add_heading(doc, "Greenness-productivity mismatch and integrated hotspots", 2)
    add_paragraph(
        doc,
        f"Hidden productivity-loss mismatch was widespread among event-affected forest pixels. "
        f"Across {int(mismatch_all['n_pixels']):,} event-affected pixels, {float(mismatch_all['hidden_mismatch_prevalence']):.2%} showed at least one year in which kNDVI had recovered while GPP remained suppressed. "
        f"The spatial-block bootstrap 95% confidence interval for this prevalence was {float(mismatch_prev_ci['ci_low']):.1%} to {float(mismatch_prev_ci['ci_high']):.1%}, deliberately wider than a pixel-level interval because nearby 0.1 degree pixels are spatially autocorrelated. "
        f"Mean hidden-mismatch duration was {float(mismatch_all['hidden_mismatch_duration_mean']):.2f} years within the t+1 to t+4 post-event window (Supplementary Table S5). "
        "The full spatial-block uncertainty table for legacy, mismatch, and recurrent-event descriptive metrics is reported in Supplementary Table S22. "
        f"The circular-shift null model produced lower hidden-mismatch prevalence ({float(null_mismatch_prev['null_mean']):.1%}) than observed ({float(null_mismatch_prev['observed_mean']):.1%}), corresponding to an observed/null ratio of {float(null_mismatch_prev['observed_to_null_ratio']):.2f}. "
        f"The mismatch was strongest in DBT, where prevalence reached {float(dbt_mismatch['hidden_mismatch_prevalence']):.2%} and mean hidden-mismatch duration reached {float(dbt_mismatch['hidden_mismatch_duration_mean']):.2f} years. "
        "This class pattern indicates that broadleaf deciduous systems frequently regained spectral greenness before regaining equivalent productivity, which is consistent with delayed carbon-uptake recovery after water and heat stress. "
        "Figure 4 is therefore interpreted as the spatial test of whether greenness recovery and productivity recovery decouple in coherent regions, while Supplementary Fig. S3 provides the t+2 anomaly diagnostic.",
    )
    doc.add_picture(str(spatial_fig3), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(
        doc,
        "Figure 4. Spatial distribution of hidden productivity-loss mismatch after compound hot-dry events. "
        "Panels show hidden mismatch duration (A), standardized kNDVI-GPP mismatch intensity (B), the fraction of GPP suppression years hidden by non-negative kNDVI anomalies (C), and GPP suppression duration (D).",
    )
    add_paragraph(
        doc,
        f"The integrated hotspot classification identified {int(severe_hotspot['n_pixels']):,} severe multi-metric hotspot pixels, covering approximately {float(severe_hotspot['area_km2']):,.0f} km2 "
        f"or {float(severe_hotspot['area_percent']):.1f}% of event-affected forest area. "
        f"The mean composite vulnerability index across event-affected forest pixels was {float(vulnerability_all['vulnerability_index_mean']):.3f}, with an 80th percentile of {float(vulnerability_all['vulnerability_index_p80']):.3f}. "
        f"DBT had the highest mean vulnerability index ({float(dbt_vulnerability['vulnerability_index_mean']):.3f}) and the largest severe-hotspot share ({float(dbt_vulnerability['severe_hotspot_percent']):.1f}%), indicating that broadleaf deciduous systems were disproportionately represented in the most vulnerable class. "
        "Because the severe class required multiple high-impact components, these hotspots are interpreted as locations where exposure, productivity loss, recurrent-event sensitivity, and greenness-productivity decoupling converged rather than as pixels selected by a single threshold. "
        "The categorical hotspot map and continuous index in Figure 5 provide the main spatial synthesis, while Supplementary Figs. S7-S8 and Supplementary Tables S7-S8 provide the vegetation-class composition and numeric class summaries.",
    )
    doc.add_picture(str(spatial_fig4), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(
        doc,
        "Figure 5. Integrated vulnerability hotspots after compound hot-dry events. "
        "Panels show the categorical hotspot class map (A), continuous vulnerability index (B), severe-hotspot domain (C), and vegetation-class context (D).",
    )
    add_heading(doc, "Spatial controls, model validation, and robustness", 2)
    add_paragraph(
        doc,
        f"The driver models indicated that GPP legacy loss and severe hotspot occurrence were predictable from event-window controls. "
        f"The random-forest model explained GPP cumulative loss with held-out R2 = {float(gpp_perf['r2']):.2f} and mean absolute error = {float(gpp_perf['mae']):.1f}. "
        f"Severe-hotspot classification reached held-out AUC = {float(hotspot_perf['auc']):.2f}. "
        f"Under spatial block cross-validation, model skill declined as expected but remained positive for the main driver outcomes, with mean GPP-loss R2 = {float(gpp_spatial['r2_mean']):.2f} across five held-out block folds and severe-hotspot AUC = {float(hotspot_spatial['auc_mean']):.2f}. "
        f"The leading permutation driver for GPP cumulative loss was {str(top_gpp_driver['feature']).replace('_', ' ')}, whereas the leading driver for severe-hotspot probability was {str(top_hotspot_driver['feature']).replace('_', ' ')}. "
        f"After GPP loss was normalized by five pre-event baseline years, the leading driver shifted to {str(top_fractional_driver['feature']).replace('_', ' ')}, while pre-event GPP importance declined to {float(fractional_baseline['importance_mean']):.2f}; spatial-block validation for the fractional-loss model remained positive but weaker (mean R2 = {float(fractional_spatial['r2_mean']):.2f}; Supplementary Fig. S14 and Supplementary Tables S23-S24). "
        "This indicates that the dominance of pre-event GPP in the absolute-loss model partly reflected the greater amount of carbon available to lose in high-productivity pixels, whereas the normalized model emphasized event heat-water stress and vegetation context. "
        "The difference between random and spatial-block validation indicates that spatial structure contributes to apparent model skill, but the blocked results show that the driver relationships are not purely local memorization. "
        "Together, Figure 6 and Figure 7 show that vulnerability was not explained by event occurrence alone; it also depended on pre-event productivity state and the intensity of heat-water stress during events. "
        "The Spearman screening, spatial-block fold metrics, logistic odds ratios, and separate driver-importance plots are reported in Supplementary Tables S9-S12 and Supplementary Figs. S9-S11.",
    )
    doc.add_picture(str(spatial_fig5), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(
        doc,
        "Figure 6. Spatial controls and model-predicted vulnerability. "
        "Panels show mean event SPEI-12 (A), mean event Tmax z-score (B), random-forest predicted vulnerability index (C), and random-forest severe-hotspot probability (D).",
    )
    doc.add_picture(str(driver_fig), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(
        doc,
        "Figure 7. Driver importance for GPP cumulative loss and severe hotspot occurrence. "
        "Permutation importance was calculated on random held-out test data; error bars show the standard deviation across permutation repeats.",
    )
    add_paragraph(
        doc,
        f"Internal robustness checks supported the main interpretation. "
        f"When NDVI replaced kNDVI in the greenness-productivity mismatch definition, {float(ndvi_robust['hidden_mismatch_prevalence']):.1%} of event-affected pixels still showed hidden productivity suppression, confirming that the mismatch was not unique to kNDVI. "
        f"Leave-one-component-out hotspot maps retained substantial overlap with the original severe-hotspot domain, with Jaccard similarity ranging from {float(robust_loo['jaccard_with_original_severe'].min()):.2f} to {float(robust_loo['jaccard_with_original_severe'].max()):.2f}. "
        f"The tree-only domain had a severe-hotspot share of {float(tree_robust['severe_hotspot_percent']):.1f}%, higher than the full event-affected vegetation domain, indicating that the hotspot signal was not driven by shrub or grass pixels. "
        "Together, these tests address three likely sources of reviewer concern: greenness-index dependence, over-reliance on one hotspot component, and sensitivity to non-tree vegetation classes. "
        "These checks are summarized in Supplementary Tables S13-S15 and Supplementary Fig. S12.",
    )
    doc.add_picture(str(robustness_fig), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(
        doc,
        "Figure 8. Internal robustness checks. "
        "Panel A compares hidden productivity-loss mismatch prevalence when kNDVI or NDVI is used as the greenness metric; panel B shows severe-hotspot stability when individual hotspot-index components are omitted.",
    )
    add_paragraph(
        doc,
        "The independent GOSIF validation directly addressed whether the hidden productivity-loss result depended on the MOD17 GPP light-use-efficiency algorithm. "
        f"When GOSIF replaced MOD17 GPP as the productivity proxy, hidden kNDVI-productivity mismatch still occurred across {int(sif_summary_all['n_pixels']):,} event-affected pixels, with prevalence of {float(sif_summary_all['hidden_mismatch_prevalence']):.1%}. "
        f"This prevalence was lower than the MOD17 GPP-based estimate ({float(sif_all['GPP_hidden_mismatch_prevalence']):.1%} over the common valid domain), indicating that MOD17 likely amplifies part of the apparent productivity suppression. "
        f"However, {float(sif_all['GPP_given_GOSIF_mismatch']):.1%} of GOSIF mismatch pixels also showed MOD17 GPP mismatch, and {float(sif_all['both_GPP_and_GOSIF_mismatch']):.1%} of all common-domain pixels showed mismatch in both productivity proxies. "
        "The persistence of the mismatch in SIF supports the interpretation that greenness recovery can coincide with incomplete functional recovery, while the weaker SIF magnitude appropriately tempers the strength of the MOD17-only carbon-loss claim.",
    )
    doc.add_picture(str(sif_fig), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(
        doc,
        "Figure 9. Independent SIF validation of hidden productivity-loss mismatch. "
        "Panel A compares all-domain mismatch prevalence from MOD17 GPP and GOSIF; panel B compares the same prevalence by vegetation class.",
    )
    add_paragraph(
        doc,
        "The greenness-side robustness test further showed that the SIF-supported mismatch was not unique to the kNDVI transformation. "
        f"When EVI replaced kNDVI, EVI recovery coincided with continued GOSIF suppression across {float(evi_structural['hidden_mismatch_prevalence']):.1%} of event-affected pixels. "
        f"When NIRv replaced kNDVI, the corresponding prevalence was {float(nirv_structural['hidden_mismatch_prevalence']):.1%}. "
        "Both estimates were lower than the kNDVI-GOSIF mismatch prevalence, but their persistence demonstrates that the result is not simply an artifact of using kNDVI, or of comparing a deterministic transform of NDVI against MOD17 GPP. "
        "NIRv provides the most conservative greenness-side test because it incorporates near-infrared canopy scattering and is more closely linked to photosynthetic canopy structure than NDVI alone.",
    )
    doc.add_picture(str(structural_fig), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(
        doc,
        "Figure 10. Structural greenness validation of hidden productivity-loss mismatch using GOSIF. "
        "Panel A compares all-domain mismatch prevalence using kNDVI, EVI, and NIRv as greenness-side recovery proxies; panel B compares the same metrics by vegetation class.",
    )
    add_paragraph(
        doc,
        f"External disturbance screening reduced the magnitude of the strongest legacy signal but did not remove it. "
        f"After excluding pixels flagged by either MODIS burned area or the Hansen forest-loss screen, {int(undisturbed['n_pixels']):,} event-affected pixels remained. "
        f"These undisturbed-screen pixels still had mean GPP cumulative loss of {float(undisturbed['GPP_cumulative_loss_mean']):.1f}, hidden mismatch prevalence of {float(undisturbed['hidden_mismatch_prevalence']):.1%}, and severe-hotspot share of {float(undisturbed['severe_hotspot_percent']):.1f}%. "
        "The disturbed domain showed stronger losses, indicating that disturbance amplifies the spatial signal, but the persistence of loss and mismatch in the undisturbed-screen domain supports a compound-climate legacy interpretation. "
        f"The severe-hotspot class did not provide external validation against the coarse Hansen forest-loss screen: forest-loss-screen prevalence was {float(hansen_severe['hansen_forest_loss_screen_prevalence']):.1%} in severe hotspots and {float(hansen_nonsevere['hansen_forest_loss_screen_prevalence']):.1%} in non-severe event pixels, and the predictive model reached only AUC = {float(hansen_model['rf_auc']):.2f} (Supplementary Table S25). "
        "Therefore, the vulnerability index is retained as an internally defined climate-response synthesis rather than presented as an externally validated mortality or forest-loss predictor.",
    )
    doc.add_picture(str(disturbance_fig), width=Inches(6.7))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(
        doc,
        "Figure 11. External disturbance-control summary. "
        "Bars compare all event pixels, pixels passing the undisturbed screen, and pixels flagged by burned-area or forest-loss disturbance for GPP cumulative loss (A), hidden mismatch prevalence (B), and severe-hotspot share (C).",
    )
    path = OUT / "Paper2_Main_Manuscript_Working_Draft.docx"
    return save_doc(doc, path)


def build_supplement():
    inventory = pd.read_csv(OBJ1 / "data" / "qgis_ready" / "objective1_qgis_raster_inventory.csv")
    doc = base_doc(
        "Supplementary Material",
        "Paper 2 working supplement",
    )
    add_heading(doc, "Supplementary Figures", 1)
    add_paragraph(
        doc,
        "This supplementary file stores supporting spatial panels, diagnostic plots, and tabular outputs for the compound hot-dry legacy analysis. "
        "Spatial panels use the common 0.1 degree EPSG:4326 grid and the study-area boundary overlay.",
    )
    supporting_figures = [
        (
            "S1",
            PUBFIG / "qgis_guide_supp_legacy_effects.png",
            "Spatial legacy effects after compound hot-dry events. Panels show kNDVI cumulative loss (A), GPP cumulative loss (B), kNDVI legacy duration (C), and GPP legacy duration (D).",
        ),
        (
            "S2",
            PUBFIG / "qgis_guide_supp_resistance_recovery_support.png",
            "Supporting recurrent-event resistance and recovery maps. Panels show GPP resistance change (A), kNDVI cumulative loss change (B), and kNDVI recovery-time change (C).",
        ),
        (
            "S3",
            PUBFIG / "figure5_kNDVI_GPP_mismatch_scatter_t2_ggstyle.png",
            "kNDVI and GPP post-event anomaly relationship at t+2. Points are sampled event-affected forest pixels and colors indicate vegetation class.",
        ),
        (
            "S4",
            PUBFIG / "figure2_first_later_loss_ggstyle.png",
            "Mean cumulative loss after first and later compound hot-dry events at recurrent-event pixels.",
        ),
        (
            "S5",
            PUBFIG / "figure3_classwise_gpp_delta_ggstyle.png",
            "Later-minus-first GPP cumulative loss by vegetation class with bootstrap 95% confidence intervals.",
        ),
        (
            "S6",
            PUBFIG / "figure4_hidden_mismatch_by_class_ggstyle.png",
            "Prevalence of hidden productivity-loss mismatch by vegetation class with bootstrap 95% confidence intervals.",
        ),
        (
            "S7",
            PUBFIG / "figure6_hotspot_composition_by_class.png",
            "Area share of integrated vulnerability-hotspot classes by vegetation class.",
        ),
        (
            "S8",
            PUBFIG / "figure7_vulnerability_index_by_class.png",
            "Mean composite vulnerability index by vegetation class.",
        ),
        (
            "S9",
            PUBFIG / "figure8_driver_importance_gpp_loss.png",
            "Random-forest permutation importance for GPP cumulative loss.",
        ),
        (
            "S10",
            PUBFIG / "figure9_driver_importance_severe_hotspot.png",
            "Random-forest permutation importance for severe-hotspot classification.",
        ),
        (
            "S11",
            PUBFIG / "figure10_heat_drought_gpp_loss_gradient.png",
            "Binned heat-drought gradient showing mean GPP cumulative loss across event-affected forest pixels.",
        ),
        (
            "S12",
            PUBFIG / "qgis_guide_supp_robustness_spatial_checks.png",
            "Spatial robustness checks. Panels show NDVI-GPP hidden mismatch duration (A), severe hotspots recalculated without event frequency (B), without hidden mismatch duration (C), and without GPP cumulative loss (D).",
        ),
        (
            "S13",
            PUBFIG / "figure16_event_timing_null_model.png",
            "Circular-shift event-timing null model for one-sided loss, legacy-duration, and hidden-mismatch metrics. Bars show observed values relative to the timing-randomized null expectation; the dashed line marks observed/null = 1.",
        ),
        (
            "S14",
            PUBFIG / "figure17_supervisor_robustness_checks.png",
            "Additional supervisor-driven robustness checks. Panel A shows random-forest permutation importance after normalizing GPP cumulative loss by pre-event baseline productivity; panel B compares compound, drought-only, and heat-only event responses; panel C compares Hansen forest-loss-screen prevalence between severe and non-severe event pixels.",
        ),
    ]
    for code, img, caption in supporting_figures:
        doc.add_picture(str(img), width=Inches(6.4))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, f"Supplementary Figure {code}. {caption}")

    add_heading(doc, "Supplementary Table S1", 1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Raster", "Bands", "CRS", "Description"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                set_run(run, size=8.5, bold=True)
    for _, row in inventory.iterrows():
        cells = table.add_row().cells
        values = [
            Path(row["file"]).name,
            str(row["bands"]),
            str(row["crs"]),
            str(row["band_descriptions"])[:130],
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run(run, size=8)
    add_caption(doc, "Supplementary Table S1. GeoTIFF outputs prepared for QGIS and manuscript mapping.")
    add_heading(doc, "Supplementary Table S2", 1)
    order = pd.read_csv(OBJ2 / "outputs" / "tables" / "event_order_summary_by_class.csv")
    add_event_order_table(doc, order)
    add_caption(doc, "Supplementary Table S2. Recurrent-event first-versus-later summary by vegetation class.")
    add_heading(doc, "Supplementary Table S3", 1)
    ci = pd.read_csv(OBJ2 / "outputs" / "tables" / "event_order_bootstrap_ci.csv")
    add_uncertainty_table(doc, ci)
    add_caption(doc, "Supplementary Table S3. Pixel bootstrap 95% confidence intervals for recurrent-event later-minus-first metrics, retained as a descriptive sensitivity check.")
    add_heading(doc, "Supplementary Table S4", 1)
    sensitivity = pd.read_csv(OBJ2 / "outputs" / "tables" / "event_order_threshold_sensitivity.csv")
    add_sensitivity_table(doc, sensitivity)
    add_caption(doc, "Supplementary Table S4. Threshold sensitivity for recurrent-event cumulative-loss results.")
    add_heading(doc, "Supplementary Table S5", 1)
    mismatch = pd.read_csv(OBJ4 / "outputs" / "tables" / "mismatch_summary_by_class.csv")
    add_mismatch_table(doc, mismatch)
    add_caption(doc, "Supplementary Table S5. Hidden productivity-loss mismatch summary by vegetation class.")
    add_heading(doc, "Supplementary Table S6", 1)
    mismatch_ci = pd.read_csv(OBJ4 / "outputs" / "tables" / "mismatch_bootstrap_ci.csv")
    add_mismatch_ci_table(doc, mismatch_ci)
    add_caption(doc, "Supplementary Table S6. Pixel bootstrap uncertainty for mismatch metrics across all event-affected forest pixels, retained as a descriptive sensitivity check.")
    add_heading(doc, "Supplementary Table S7", 1)
    hotspot = pd.read_csv(OBJ3 / "outputs" / "tables" / "hotspot_class_summary_by_vegetation.csv")
    add_hotspot_summary_table(doc, hotspot)
    add_caption(doc, "Supplementary Table S7. Integrated hotspot-class summary across all event-affected forest pixels.")
    add_heading(doc, "Supplementary Table S8", 1)
    vulnerability = pd.read_csv(OBJ3 / "outputs" / "tables" / "vulnerability_index_summary_by_vegetation.csv")
    add_vulnerability_table(doc, vulnerability)
    add_caption(doc, "Supplementary Table S8. Composite vulnerability index and severe-hotspot share by vegetation class.")
    add_heading(doc, "Supplementary Table S9", 1)
    model_perf = pd.read_csv(OBJ5 / "outputs" / "tables" / "model_performance_summary.csv")
    add_model_performance_table(doc, model_perf)
    add_caption(doc, "Supplementary Table S9. Held-out model performance for driver/control models.")
    add_heading(doc, "Supplementary Table S9b", 1)
    spatial_cv_folds = pd.read_csv(OBJ5 / "outputs" / "tables" / "spatial_block_cv_fold_metrics.csv")
    add_spatial_cv_fold_table(doc, spatial_cv_folds)
    add_caption(
        doc,
        "Supplementary Table S9b. Five-fold spatial block cross-validation metrics for driver/control models. "
        "Each fold holds out independent 10 degree geographic blocks.",
    )
    add_heading(doc, "Supplementary Table S10", 1)
    gpp_importance = pd.read_csv(OBJ5 / "outputs" / "tables" / "rf_importance_GPP_cumulative_loss.csv")
    severe_importance = pd.read_csv(OBJ5 / "outputs" / "tables" / "rf_importance_severe_hotspot.csv")
    add_importance_table(doc, gpp_importance, severe_importance)
    add_caption(doc, "Supplementary Table S10. Leading random-forest driver importances.")
    add_heading(doc, "Supplementary Table S11", 1)
    corr = pd.read_csv(OBJ5 / "outputs" / "tables" / "driver_spearman_correlations.csv")
    top_corr = corr.reindex(corr["spearman_rho"].abs().sort_values(ascending=False).index).head(12)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Outcome", "Predictor", "Spearman rho", "n"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.2, bold=True)
    for _, row in top_corr.iterrows():
        cells = table.add_row().cells
        values = [str(row["outcome"]), str(row["predictor_label"]), f"{float(row['spearman_rho']):.3f}", f"{int(row['n']):,}"]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=8)
    add_caption(doc, "Supplementary Table S11. Strongest absolute Spearman driver correlations across outcomes.")
    add_heading(doc, "Supplementary Table S12", 1)
    odds = pd.read_csv(OBJ5 / "outputs" / "tables" / "logistic_odds_ratios_severe_hotspot.csv").head(12)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Feature", "Odds ratio", "95% CI low", "95% CI high"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.2, bold=True)
    for _, row in odds.iterrows():
        cells = table.add_row().cells
        values = [str(row["feature"]), f"{float(row['odds_ratio']):.2f}", f"{float(row['ci_low']):.2f}", f"{float(row['ci_high']):.2f}"]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=8)
    add_caption(doc, "Supplementary Table S12. Logistic-regression odds ratios for severe-hotspot occurrence using standardized continuous predictors.")
    add_heading(doc, "Supplementary Table S13", 1)
    robust_greenness = pd.read_csv(OBJ99 / "outputs" / "tables" / "greenness_index_mismatch_robustness.csv")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Index", "Class", "Pixels", "Prevalence", "Duration"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.2, bold=True)
    for _, row in robust_greenness.iterrows():
        cells = table.add_row().cells
        values = [
            str(row["greenness_index"]),
            str(row["vegetation_class"]),
            f"{int(row['n_pixels']):,}",
            f"{float(row['hidden_mismatch_prevalence']):.3f}",
            f"{float(row['hidden_mismatch_duration_mean']):.2f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=8)
    add_caption(doc, "Supplementary Table S13. Hidden productivity-loss mismatch robustness when NDVI is substituted for kNDVI.")
    add_heading(doc, "Supplementary Table S14", 1)
    robust_loo = pd.read_csv(OBJ99 / "outputs" / "tables" / "hotspot_leave_one_component_out_stability.csv")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Omitted component", "Severe pixels", "Jaccard", "Pixel agreement"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.2, bold=True)
    for _, row in robust_loo.iterrows():
        cells = table.add_row().cells
        values = [
            str(row["omitted_component"]),
            f"{int(row['severe_pixels']):,}",
            f"{float(row['jaccard_with_original_severe']):.3f}",
            f"{float(row['pixel_agreement_with_original']):.3f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=8)
    add_caption(doc, "Supplementary Table S14. Leave-one-component-out stability of severe-hotspot classification.")
    add_heading(doc, "Supplementary Table S15", 1)
    robust_tree = pd.read_csv(OBJ99 / "outputs" / "tables" / "tree_only_domain_robustness.csv")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Domain", "Pixels", "Events", "Mean vulnerability", "Severe hotspots"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=8.2, bold=True)
    for _, row in robust_tree.iterrows():
        cells = table.add_row().cells
        values = [
            str(row["domain"]),
            f"{int(row['n_pixels']):,}",
            f"{float(row['mean_event_frequency']):.2f}",
            f"{float(row['mean_vulnerability_index']):.3f}",
            f"{float(row['severe_hotspot_percent']):.1f}%",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=8)
    add_caption(doc, "Supplementary Table S15. Robustness of hotspot metrics under a tree-only domain restriction.")
    add_heading(doc, "Supplementary Table S16", 1)
    disturbance = pd.read_csv(OBJ99 / "outputs" / "tables" / "external_disturbance_control_summary.csv")
    add_external_disturbance_table(doc, disturbance)
    add_caption(doc, "Supplementary Table S16. External disturbance-control summary using MODIS burned area and Hansen forest-loss screening.")
    add_heading(doc, "Supplementary Table S17", 1)
    sif_summary = pd.read_csv(OBJ6 / "outputs" / "tables" / "gosif_mismatch_summary_by_class.csv")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Class", "Pixels", "Event count", "GOSIF mismatch duration", "GOSIF mismatch prevalence", "GOSIF suppression duration"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=7.7, bold=True)
    for _, row in sif_summary.iterrows():
        cells = table.add_row().cells
        values = [
            str(row["vegetation_class"]),
            f"{int(row['n_pixels']):,}",
            f"{float(row['event_count_mean']):.2f}",
            f"{float(row['hidden_mismatch_duration_mean']):.2f}",
            f"{float(row['hidden_mismatch_prevalence']):.3f}",
            f"{float(row['productivity_suppression_duration_mean']):.2f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=7.2)
    add_caption(doc, "Supplementary Table S17. Hidden productivity-loss mismatch recalculated using GOSIF as an independent productivity proxy.")
    add_heading(doc, "Supplementary Table S18", 1)
    sif_comparison = pd.read_csv(OBJ6 / "outputs" / "tables" / "gpp_vs_gosif_mismatch_comparison.csv")
    add_sif_validation_table(doc, sif_comparison)
    add_caption(doc, "Supplementary Table S18. Common-domain overlap between MOD17 GPP-based and GOSIF-based hidden mismatch.")
    add_heading(doc, "Supplementary Table S19", 1)
    structural_greenness = pd.read_csv(OBJ6 / "outputs" / "tables" / "structural_greenness_gosif_validation.csv")
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["Greenness", "Productivity", "Class", "Pixels", "Prevalence", "Duration", "GOSIF suppression"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                set_run(run, size=7.6, bold=True)
    for _, row in structural_greenness.iterrows():
        cells = table.add_row().cells
        values = [
            str(row["greenness_proxy"]),
            str(row["productivity_proxy"]),
            str(row["vegetation_class"]),
            f"{int(row['n_pixels']):,}",
            f"{float(row['hidden_mismatch_prevalence']):.3f}",
            f"{float(row['hidden_mismatch_duration_mean']):.2f}",
            f"{float(row['GOSIF_suppression_duration_mean']):.2f}",
        ]
        for cell, value in zip(cells, values):
            cell.text = value
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    set_run(run, size=7.0)
    add_caption(doc, "Supplementary Table S19. Structural greenness validation using kNDVI, EVI, and NIRv against GOSIF suppression.")
    add_heading(doc, "Supplementary Table S20", 1)
    null_summary = pd.read_csv(OBJ99 / "outputs" / "tables" / "event_timing_null_summary.csv")
    add_event_timing_null_table(doc, null_summary)
    add_caption(doc, "Supplementary Table S20. Circular-shift event-timing null model for one-sided loss, legacy-duration, and hidden-mismatch metrics across all event-affected pixels.")
    add_heading(doc, "Supplementary Table S21", 1)
    clean_gap = pd.read_csv(OBJ2 / "outputs" / "tables" / "clean_gap_event_order_summary.csv")
    add_clean_gap_table(doc, clean_gap)
    add_caption(doc, "Supplementary Table S21. Clean-gap recurrent-event sensitivity requiring at least eight years between the previous valid event and the later event, preventing baseline-response window overlap.")
    add_heading(doc, "Supplementary Table S22", 1)
    spatial_desc = pd.read_csv(OBJ99 / "outputs" / "tables" / "spatial_block_descriptive_bootstrap.csv")
    add_spatial_block_bootstrap_table(doc, spatial_desc)
    add_caption(doc, "Supplementary Table S22. Spatial-block bootstrap uncertainty for key descriptive statistics using 10 degree geographic blocks resampled with replacement.")
    add_heading(doc, "Supplementary Table S23", 1)
    fractional_imp = pd.read_csv(OBJ99 / "outputs" / "tables" / "rf_importance_GPP_fractional_cumulative_loss.csv")
    add_simple_dataframe_table(
        doc,
        fractional_imp,
        ["feature", "importance_mean", "importance_sd"],
        ["Feature", "Importance", "SD"],
        max_rows=12,
    )
    add_caption(doc, "Supplementary Table S23a. Random-forest permutation importance after normalizing GPP cumulative loss by five pre-event baseline years.")
    fractional_perf = pd.read_csv(OBJ99 / "outputs" / "tables" / "fractional_loss_model_performance.csv")
    add_simple_dataframe_table(
        doc,
        fractional_perf,
        ["model", "n_test", "r2", "mae", "n_folds", "r2_mean", "r2_sd", "mae_mean"],
        ["Model", "Test n", "R2", "MAE", "Folds", "Block R2", "Block R2 SD", "Block MAE"],
    )
    add_caption(doc, "Supplementary Table S23b. Held-out and spatial-block validation performance for the fractional GPP cumulative-loss model.")
    add_heading(doc, "Supplementary Table S24", 1)
    event_response = pd.read_csv(OBJ99 / "outputs" / "tables" / "compound_vs_univariate_event_response.csv")
    add_simple_dataframe_table(
        doc,
        event_response,
        [
            "event_type",
            "n_event_pixels",
            "event_count_mean",
            "GPP_cumulative_loss_mean",
            "GPP_fractional_cumulative_loss_mean",
            "GPP_legacy_duration_mean",
            "GPP_pre_event_baseline_mean",
        ],
        ["Event type", "Pixels", "Mean events", "GPP loss", "Fractional loss", "GPP duration", "Pre-event GPP"],
    )
    add_caption(doc, "Supplementary Table S24. Mutually exclusive compound, drought-only, and heat-only event response comparison using identical baseline and response windows.")
    add_heading(doc, "Supplementary Table S25", 1)
    hansen_validation = pd.read_csv(OBJ99 / "outputs" / "tables" / "hotspot_hansen_forest_loss_validation.csv")
    add_simple_dataframe_table(
        doc,
        hansen_validation,
        [
            "domain",
            "n_pixels",
            "hansen_forest_loss_screen_prevalence",
            "mean_vulnerability_index",
            "burned_screen_prevalence",
            "rf_auc",
            "spearman_vulnerability_loss_rho",
        ],
        ["Domain", "Pixels", "Hansen loss screen", "Mean vulnerability", "Burned screen", "RF AUC", "Spearman rho"],
        percent_cols=["hansen_forest_loss_screen_prevalence", "burned_screen_prevalence"],
    )
    add_caption(doc, "Supplementary Table S25. External validation check comparing the severe-hotspot domain with Hansen forest-loss and MODIS burned-area screens.")
    add_heading(doc, "Supplementary Table S26", 1)
    legacy = pd.read_csv(OBJ1 / "outputs" / "tables" / "objective1_legacy_summary_by_class.csv")
    add_key_table(doc, legacy)
    add_caption(
        doc,
        "Supplementary Table S26. Legacy metrics by vegetation class. "
        "Vegetation classes are evergreen needleleaf trees (ENT), evergreen broadleaf trees (EBT), deciduous needleleaf trees (DNT), deciduous broadleaf trees (DBT), shrubs (SHB), and grasslands (GRS).",
    )
    path = OUT / "Paper2_Supplementary_Material_Working_Draft.docx"
    return save_doc(doc, path)


def main():
    main_path = build_main()
    supp_path = build_supplement()
    print(main_path)
    print(supp_path)


if __name__ == "__main__":
    main()
